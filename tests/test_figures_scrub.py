"""Figure metadata stripping (METADATA_PRIVACY item 13).

The project tree LaTeXtify writes is itself a deliverable -- ``figures/`` goes
to arXiv and to journal source-file uploads -- so a photo that reaches it as a
passthrough copy must not still carry the camera's GPS fix, the body/lens
serial numbers, the photographer's name, or the capture thumbnail (which can
show an *uncropped* original frame).

The stripping is required to be lossless, which is what separates it from
:func:`latextify.privacy.images.sanitize`: that function rebuilds from pixel
data and re-encodes, MEASURED at a worst per-channel delta of 64/255 on a
quality-95 JPEG. These tests therefore assert on the pixels and the ICC
profile as hard as they assert on the metadata -- a privacy feature that
degrades every figure in every conversion is not shippable.

The leaky images are generated here at test run time with Pillow; the shared
``tests/fixtures/leaky_photo.jpg`` (with its ``.truth.json`` sidecar) is used
for the end-to-end manuscript case so the planted leaks are the same ones the
privacy suite asserts on.
"""

from __future__ import annotations

import io
import shutil
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageCms, PngImagePlugin

from latextify.emit.project import emit_project
from latextify.figures.scrub import strip_figure_metadata
from latextify.privacy import images

FIXTURES = Path(__file__).parent / "fixtures"
LEAKY_PHOTO = FIXTURES / "leaky_photo.jpg"


def _srgb_profile() -> bytes:
    return ImageCms.ImageCmsProfile(ImageCms.createProfile("sRGB")).tobytes()


def _gradient(size=(64, 48)) -> Image.Image:
    """A non-uniform image, so a JPEG re-encode would actually show as a delta.

    A flat colour block survives requantization almost unchanged, which would
    make the losslessness assertions pass for the wrong reason.
    """
    width, height = size
    image = Image.new("RGB", size)
    pixels = image.load()
    for y in range(height):
        for x in range(width):
            pixels[x, y] = ((x * 7) % 256, (y * 5) % 256, ((x + y) * 3) % 256)
    return image


def _leaky_jpeg(path: Path, *, icc: bytes | None = None) -> Image.Image:
    """Write a JPEG carrying EXIF (make/model/artist) and GPS; return the source image."""
    image = _gradient()
    exif = image.getexif()
    exif[271] = "Testcorp"  # Make
    exif[272] = "Microscope 9000"  # Model
    exif[315] = "Dr Testy McTestface"  # Artist
    gps = exif.get_ifd(0x8825)
    gps[1] = "N"
    gps[2] = (39.0, 8.0, 17.48)
    gps[3] = "W"
    gps[4] = (77.0, 13.0, 11.36)
    image.save(path, quality=95, exif=exif, **({"icc_profile": icc} if icc else {}))
    return image


def _leaky_png(path: Path, *, icc: bytes | None = None) -> Image.Image:
    image = _gradient()
    info = PngImagePlugin.PngInfo()
    info.add_text("Author", "Dr Testy McTestface")
    info.add_text("Software", "SecretLabSuite 3.1")
    exif = image.getexif()
    exif[315] = "Dr Testy McTestface"
    image.save(path, pnginfo=info, exif=exif, **({"icc_profile": icc} if icc else {}))
    return image


def _png_chunks(path: Path) -> list[str]:
    data = path.read_bytes()
    offset, found = 8, []
    while offset + 8 <= len(data):
        length = int.from_bytes(data[offset : offset + 4], "big")
        found.append(data[offset + 4 : offset + 8].decode("latin1"))
        offset += 12 + length
    return found


# --------------------------------------------------------------------------- #
# the strip itself: metadata gone, pixels and colour profile untouched
# --------------------------------------------------------------------------- #


def test_jpeg_loses_its_exif_and_gps(tmp_path):
    figure = tmp_path / "fig1.jpg"
    _leaky_jpeg(figure)
    assert images.inspect(figure)[0], "fixture should start out leaky"

    assert strip_figure_metadata(figure) is None
    assert images.inspect(figure)[0] == []


def test_jpeg_pixels_are_bit_identical_after_stripping(tmp_path):
    """The whole reason this does not reuse privacy.images.sanitize."""
    figure = tmp_path / "fig1.jpg"
    _leaky_jpeg(figure)
    before = Image.open(figure).tobytes()

    strip_figure_metadata(figure)

    after = Image.open(figure)
    assert after.tobytes() == before
    assert after.size == (64, 48)


def test_jpeg_keeps_its_icc_profile(tmp_path):
    figure = tmp_path / "fig1.jpg"
    _leaky_jpeg(figure, icc=_srgb_profile())

    strip_figure_metadata(figure)

    assert Image.open(figure).info.get("icc_profile")
    assert images.inspect(figure)[0] == []


def _jpeg_markers(path: Path) -> list[int]:
    """The marker byte of every segment up to (and including) the scan header."""
    data, offset, found = path.read_bytes(), 2, []
    while offset + 1 < len(data):
        marker = data[offset + 1]
        found.append(marker)
        if marker == 0xDA:  # SOS -- entropy data follows, stop walking
            break
        offset += 2 + int.from_bytes(data[offset + 2 : offset + 4], "big")
    return found


def _insert_app1(path: Path, payload: bytes) -> None:
    """Splice a raw APP1 segment in after SOI, the way a camera writes EXIF."""
    data = path.read_bytes()
    segment = b"\xff\xe1" + (len(payload) + 2).to_bytes(2, "big") + payload
    path.write_bytes(data[:2] + segment + data[2:])


def test_jpeg_thumbnail_does_not_survive(tmp_path):
    """An EXIF thumbnail is not regenerated on edit, so it can still show the
    original uncropped frame. It rides in APP1, so the whole segment must go."""
    figure = tmp_path / "fig1.jpg"
    _gradient().save(figure, quality=95)
    thumbnail = io.BytesIO()
    _gradient(size=(8, 6)).save(thumbnail, format="JPEG")
    _insert_app1(figure, b"Exif\x00\x00" + thumbnail.getvalue())
    # The fixture must actually carry the thing this test is about.
    assert 0xE1 in _jpeg_markers(figure)
    assert thumbnail.getvalue() in figure.read_bytes()

    assert strip_figure_metadata(figure) is None

    assert 0xE1 not in _jpeg_markers(figure)
    assert thumbnail.getvalue() not in figure.read_bytes()


def test_jpeg_keeps_the_segments_a_decoder_needs(tmp_path):
    """Dropping APP0/APP14 or the quantization/frame headers would break rendering."""
    figure = tmp_path / "fig1.jpg"
    _leaky_jpeg(figure)
    before = set(_jpeg_markers(figure))

    strip_figure_metadata(figure)

    kept = set(_jpeg_markers(figure))
    assert 0xE1 not in kept  # APP1 (EXIF/GPS) gone
    assert before - kept == {0xE1}  # and nothing else was touched


def test_app2_is_kept_only_when_it_is_an_icc_profile(tmp_path):
    """APP2 also carries MPF, which on phone cameras embeds a whole second JPEG."""
    figure = tmp_path / "fig1.jpg"
    _gradient().save(figure, quality=95)
    data = figure.read_bytes()
    mpf = b"MPF\x00" + b"embedded-original-frame"
    segment = b"\xff\xe2" + (len(mpf) + 2).to_bytes(2, "big") + mpf
    figure.write_bytes(data[:2] + segment + data[2:])
    assert 0xE2 in _jpeg_markers(figure)

    strip_figure_metadata(figure)

    assert 0xE2 not in _jpeg_markers(figure)
    assert b"embedded-original-frame" not in figure.read_bytes()


def test_a_file_that_is_not_really_an_image_is_left_silently_alone(tmp_path):
    """No PNG signature means no PNG metadata, so a privacy warning would be a
    false positive; the compile step reports the bad figure far better."""
    figure = tmp_path / "fig1.png"
    figure.write_bytes(b"not really a png")

    assert strip_figure_metadata(figure) is None
    assert figure.read_bytes() == b"not really a png"


def test_png_loses_its_text_chunks_and_exif(tmp_path):
    figure = tmp_path / "fig1.png"
    _leaky_png(figure)
    assert "tEXt" in _png_chunks(figure)

    assert strip_figure_metadata(figure) is None

    assert "tEXt" not in _png_chunks(figure)
    assert "eXIf" not in _png_chunks(figure)
    assert not Image.open(figure).getexif()


def test_png_keeps_pixels_and_colour_profile(tmp_path):
    figure = tmp_path / "fig1.png"
    _leaky_png(figure, icc=_srgb_profile())
    before = Image.open(figure).tobytes()

    strip_figure_metadata(figure)

    after = Image.open(figure)
    assert after.tobytes() == before
    assert after.info.get("icc_profile")
    assert "iCCP" in _png_chunks(figure)


def test_a_clean_image_is_left_byte_identical(tmp_path):
    """Nothing to remove must mean nothing written -- not a gratuitous rewrite."""
    figure = tmp_path / "fig1.png"
    _gradient().save(figure)
    before = figure.read_bytes()

    assert strip_figure_metadata(figure) is None
    assert figure.read_bytes() == before


@pytest.mark.parametrize("name", ["fig1.pdf", "fig1.eps", "fig1.svg"])
def test_non_raster_figures_are_left_alone(tmp_path, name):
    figure = tmp_path / name
    figure.write_bytes(b"%PDF-1.4 not really")

    assert strip_figure_metadata(figure) is None
    assert figure.read_bytes() == b"%PDF-1.4 not really"


# --------------------------------------------------------------------------- #
# failure degrades to a warning; the figure is never lost or truncated
# --------------------------------------------------------------------------- #


def test_unparseable_image_warns_and_keeps_the_original(tmp_path):
    figure = tmp_path / "fig1.jpg"
    figure.write_bytes(b"\xff\xd8this is not a JPEG body")

    warning = strip_figure_metadata(figure)

    assert warning is not None
    assert "fig1.jpg" in warning
    assert "still present" in warning  # says what was NOT achieved
    assert figure.read_bytes() == b"\xff\xd8this is not a JPEG body"
    assert [p.name for p in tmp_path.iterdir()] == ["fig1.jpg"]  # no temp left behind


def test_truncated_jpeg_is_never_rewritten(tmp_path):
    """A real JPEG cut off before its scan: rewriting what parsed so far would
    replace the original with a plausible-looking truncation."""
    figure = tmp_path / "fig1.jpg"
    _leaky_jpeg(figure)
    data = figure.read_bytes()
    figure.write_bytes(data[: data.index(b"\xff\xda")])  # everything before SOS
    truncated = figure.read_bytes()

    warning = strip_figure_metadata(figure)

    assert warning is not None and "fig1.jpg" in warning
    assert figure.read_bytes() == truncated


def test_missing_file_is_not_an_error(tmp_path):
    """The TIFF path writes nothing when conversion fails; scrubbing must not care."""
    assert strip_figure_metadata(tmp_path / "never-written.png") is None


# --------------------------------------------------------------------------- #
# end to end: a manuscript's figure reaches the output tree clean
# --------------------------------------------------------------------------- #


def _manuscript_with(image_path: Path, dest: Path) -> None:
    document = Document()
    document.add_heading("Scrub Fixture", level=0)
    document.add_heading("Results", level=1)
    document.add_paragraph("A placeholder result, illustrated below.")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("Figure 1: A placeholder result.", style="Caption")
    document.save(dest)


@pytest.fixture()
def leaky_manuscript(tmp_path):
    docx_path = tmp_path / "leaky_figure.docx"
    _manuscript_with(LEAKY_PHOTO, docx_path)
    return docx_path


def test_emitted_figure_carries_no_metadata_by_default(leaky_manuscript, tmp_path):
    assert images.inspect(LEAKY_PHOTO)[0], "the shared fixture should be leaky"

    result = emit_project(leaky_manuscript, "revtex4-2", tmp_path / "out")

    emitted = result.output_dir / "figures" / "fig1.jpg"
    assert emitted.is_file()
    assert images.inspect(emitted)[0] == []


def test_keep_figure_metadata_preserves_it(leaky_manuscript, tmp_path):
    result = emit_project(
        leaky_manuscript, "revtex4-2", tmp_path / "out", strip_figure_metadata=False
    )

    emitted = result.output_dir / "figures" / "fig1.jpg"
    summaries = [f.summary for f in images.inspect(emitted)[0]]
    assert any("GPS" in s for s in summaries)
    assert any("Dr Testy McTestface" in s for s in summaries)


def test_stripping_does_not_change_the_emitted_pixels(leaky_manuscript, tmp_path):
    """A privacy default that silently degrades figures would be worse than the leak."""
    result = emit_project(leaky_manuscript, "revtex4-2", tmp_path / "out")

    emitted = Image.open(result.output_dir / "figures" / "fig1.jpg")
    original = Image.open(LEAKY_PHOTO)
    assert emitted.size == original.size
    assert emitted.tobytes() == original.tobytes()


def test_unstrippable_figure_surfaces_as_an_emit_warning(leaky_manuscript, tmp_path, monkeypatch):
    """A figure the stripper cannot parse must be reported, never silently shipped."""
    import latextify.figures.convert as convert_module

    monkeypatch.setattr(
        convert_module, "strip_figure_metadata", lambda path: "could not strip test-figure"
    )
    result = emit_project(leaky_manuscript, "revtex4-2", tmp_path / "out")

    assert any("could not strip test-figure" in w.message for w in result.warnings)
    assert (result.output_dir / "figures" / "fig1.jpg").is_file()


def test_supplement_figures_are_stripped_too(tmp_path):
    """figures/ is shared by both documents; the SI's figS<N> must not be a hole."""
    main_docx = tmp_path / "main.docx"
    si_docx = tmp_path / "si.docx"
    si_photo = tmp_path / "si_photo.jpg"
    shutil.copy2(LEAKY_PHOTO, si_photo)
    _manuscript_with(LEAKY_PHOTO, main_docx)
    _manuscript_with(si_photo, si_docx)

    result = emit_project(
        main_docx, "revtex4-2", tmp_path / "out", supplement_docx_path=si_docx
    )

    emitted = result.output_dir / "figures" / "figS1.jpg"
    assert emitted.is_file()
    assert images.inspect(emitted)[0] == []
