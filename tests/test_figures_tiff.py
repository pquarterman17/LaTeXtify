"""End-to-end TIFF figure conversion (GAP 2): a real manuscript-conversion bug.

Word manuscripts embed TIFF images constantly (scanner/microscope exports
commonly land in a .docx this way); Tectonic's xdvipdfmx PDF backend has no
TIFF support at all, so a raw \\includegraphics of a .tif/.tiff fails to
compile with "Cannot determine size of graphic" -- this is exactly the
failure a real manuscript conversion hit. latextify.figures.convert now
converts .tif/.tiff -> .png via Pillow at copy time (mirrors the SVG->PDF and
EPS->PDF conversion paths, plan item 15).

These tests build a synthesized (fake-content) manuscript with an embedded
TIFF image, generating the TIFF itself with Pillow entirely at test run
time -- no committed binary fixture, and nothing from any real manuscript.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from latextify.compile.tectonic import compile_document, ensure_tectonic, find_tectonic
from latextify.emit.project import emit_project


def _tiff_bytes(color: tuple[int, int, int] = (200, 30, 30), size=(16, 16)) -> io.BytesIO:
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="TIFF")
    buf.seek(0)
    return buf


def _build_tiff_docx(path) -> None:
    document = Document()
    document.add_heading("TIFF Figure Fixture", level=0)
    document.add_heading("Results", level=1)
    document.add_paragraph("A placeholder result, illustrated below.")
    document.add_picture(_tiff_bytes(), width=Inches(1))
    document.add_paragraph(
        "Figure 1: A placeholder result embedded as a TIFF image.", style="Caption"
    )
    document.save(path)


@pytest.fixture()
def tiff_docx(tmp_path):
    docx_path = tmp_path / "tiff_figure.docx"
    _build_tiff_docx(docx_path)
    return docx_path


# --------------------------------------------------------------------------- #
# embedded TIFF media -> PNG in the output tree
# --------------------------------------------------------------------------- #


def test_embedded_tiff_lands_as_png_in_output_tree(tiff_docx, tmp_path):
    result = emit_project(tiff_docx, "revtex4-2", tmp_path / "output", report=False)

    figures_written = list(result.figures_dir.iterdir())
    assert figures_written, "no figures written"
    assert all(f.suffix == ".png" for f in figures_written)
    assert not any(f.suffix in (".tif", ".tiff") for f in figures_written)

    png_path = result.figures_dir / "fig1.png"
    assert png_path.is_file()
    assert png_path.read_bytes().startswith(b"\x89PNG\r\n\x1a\n")

    # Figure IR carries the conversion note, mirroring the SVG->PDF path.
    assert result.figures[0].conversion_note is not None
    assert "TIFF" in result.figures[0].conversion_note

    # A clean Pillow conversion raises no TIFF-related warning.
    assert not any("Pillow" in w.message or "TIFF" in w.message for w in result.warnings)


def test_embedded_tiff_body_includes_the_converted_png(tiff_docx, tmp_path):
    result = emit_project(tiff_docx, "revtex4-2", tmp_path / "output", report=False)
    body = result.body_tex_path.read_text(encoding="utf-8")
    assert "\\includegraphics{figures/fig1.png}" in body
    assert ".tif" not in body


# --------------------------------------------------------------------------- #
# folder-convention override: a figures/fig1.tiff override must ALSO convert
# --------------------------------------------------------------------------- #


def test_override_tiff_also_converts_to_png(tmp_path):
    docx_path = tmp_path / "override.docx"
    _build_tiff_docx(docx_path)
    figures_dir = tmp_path / "figures"
    figures_dir.mkdir()
    Image.new("RGB", (16, 16), (10, 200, 10)).save(figures_dir / "fig1.tiff", format="TIFF")

    result = emit_project(docx_path, "revtex4-2", tmp_path / "output", report=False)

    png_path = result.figures_dir / "fig1.png"
    assert png_path.is_file()
    assert png_path.read_bytes().startswith(b"\x89PNG\r\n\x1a\n")
    assert not any(f.suffix in (".tif", ".tiff") for f in result.figures_dir.iterdir())


# --------------------------------------------------------------------------- #
# empirical: does the converted PNG actually compile under Tectonic?
# --------------------------------------------------------------------------- #


def _tectonic_available() -> bool:
    return find_tectonic() is not None


@pytest.mark.tectonic
@pytest.mark.skipif(
    not _tectonic_available(),
    reason="no tectonic binary on PATH/cache and none could be downloaded",
)
def test_tiff_figure_manuscript_compiles_to_pdf(tiff_docx, tmp_path):
    # Before the fix, this exact shape ("Cannot determine size of graphic")
    # is the real manuscript-conversion failure that motivated GAP 2.
    result = emit_project(tiff_docx, "revtex4-2", tmp_path / "output", report=False)

    compile_result = compile_document(result.main_tex_path, tectonic_path=ensure_tectonic())

    assert compile_result.success, compile_result.raw_log
    assert compile_result.pdf_path is not None
    assert compile_result.pdf_path.is_file()
    assert compile_result.pdf_path.stat().st_size > 0
