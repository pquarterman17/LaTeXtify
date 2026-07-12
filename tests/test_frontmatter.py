"""Front-matter stripping (gap 4): the manuscript's own title page must not be
duplicated in the body.

The journal metadata template re-renders title/authors/affiliations/abstract/
keywords from paper.yaml; pandoc converts the body verbatim. Without stripping,
a real manuscript's title page appears twice in the PDF. These tests cover the
span detector's conservative gate (strip only on a strong title-page signal)
and the end-to-end emit behaviour (title page gone from body.tex, still present
in metadata.tex), for several title-page shapes -- not just the one real paper
that surfaced the bug.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from latextify.ingest.frontmatter import strip_front_matter_from_docx
from latextify.ingest.metadata_guess import front_matter_span
from latextify.ingest.pandoc import convert_docx_to_body


def _add_superscript(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.superscript = True


def _build_manuscript(
    path: Path,
    *,
    markers: bool = True,
    abstract: bool = True,
    keywords: bool = True,
    corresponding: bool = True,
) -> Path:
    """A manuscript-shaped docx: title page followed by a body section.

    Flags toggle which title-page elements are present so the tests can
    exercise different shapes (no keywords, no abstract, marker-only, ...).
    """
    docx = pytest.importorskip("docx")
    doc = docx.Document()

    doc.add_paragraph(style="Title").add_run("Frustrated Coupling in Model Heterostructures")

    authors = doc.add_paragraph()
    authors.add_run("Alice Author")
    if markers:
        _add_superscript(authors, "1*" if corresponding else "1")
    authors.add_run(", Bob Builder")
    if markers:
        _add_superscript(authors, "2")

    aff1 = doc.add_paragraph()
    if markers:
        _add_superscript(aff1, "1")
    aff1.add_run("Institute of Physics, Example University, Exampletown")

    aff2 = doc.add_paragraph()
    if markers:
        _add_superscript(aff2, "2")
    aff2.add_run("Department of Materials, Sample Institute, Sampleville")

    if corresponding:
        corr = doc.add_paragraph()
        _add_superscript(corr, "*")
        corr.add_run("Corresponding author: alice@example.edu")

    if abstract:
        doc.add_paragraph("Abstract")
        doc.add_paragraph(
            "We report a study of frustrated magnetic coupling in model "
            "heterostructures and its dependence on layer composition."
        )

    if keywords:
        doc.add_paragraph("Keywords: frustration, heterostructures, coupling")

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "The study of magnetic frustration has a long history in condensed "
        "matter physics, motivating the measurements presented below."
    )

    doc.save(path)
    return path


# --------------------------------------------------------------------------- #
# Span detection + conservative gate
# --------------------------------------------------------------------------- #


def test_span_covers_full_title_page(tmp_path):
    path = _build_manuscript(tmp_path / "full.docx")
    span = front_matter_span(path)
    assert span is not None
    start, end = span
    # Title page is paragraphs 0..7 (title, authors, 2 affiliations, corr line,
    # "Abstract", abstract body, keywords); the body "Introduction" heading is
    # paragraph 8 and must NOT be included.
    assert start == 0
    assert end == 8


def test_span_none_without_markers_or_abstract(tmp_path):
    # A bare Title heading + body, no author markers and no Abstract heading:
    # too weak a signal (this is the figures-only-fixture shape). Strip nothing.
    path = _build_manuscript(
        tmp_path / "bare.docx", markers=False, abstract=False, keywords=False, corresponding=False
    )
    assert front_matter_span(path) is None


def test_span_triggered_by_markers_alone(tmp_path):
    path = _build_manuscript(
        tmp_path / "markers_only.docx", abstract=False, keywords=False, corresponding=False
    )
    span = front_matter_span(path)
    assert span is not None
    # title + author line + two affiliations = paragraphs 0..3; body heading at 4.
    assert span == (0, 4)


def test_span_triggered_by_abstract_alone(tmp_path):
    path = _build_manuscript(
        tmp_path / "abstract_only.docx", markers=False, keywords=False, corresponding=False
    )
    span = front_matter_span(path)
    assert span is not None


def test_span_handles_missing_keywords(tmp_path):
    path = _build_manuscript(tmp_path / "no_kw.docx", keywords=False)
    span = front_matter_span(path)
    assert span is not None
    # No keywords line: span ends after the abstract body, before "Introduction".
    _, end = span
    assert end == 7


# --------------------------------------------------------------------------- #
# Stripping mechanics
# --------------------------------------------------------------------------- #


def test_strip_removes_title_page_paragraphs(tmp_path):
    path = _build_manuscript(tmp_path / "full.docx")
    stripped = strip_front_matter_from_docx(path, tmp_path / "work")
    assert stripped != path  # a rewritten copy was produced

    from lxml import etree

    from latextify.citations.fields import read_document_xml

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = etree.fromstring(read_document_xml(stripped))
    body = root.find(f"{{{W}}}body")
    texts = ["".join(p.itertext()) for p in body.findall(f"{{{W}}}p")]
    joined = "\n".join(texts)

    # Title-page content gone.
    assert "Frustrated Coupling in Model Heterostructures" not in joined
    assert "Alice Author" not in joined
    assert "Institute of Physics" not in joined
    assert "We report a study of frustrated" not in joined
    # Body survives.
    assert "Introduction" in joined
    assert "magnetic frustration has a long history" in joined


def test_strip_passthrough_when_no_front_matter(tmp_path):
    path = _build_manuscript(
        tmp_path / "bare.docx", markers=False, abstract=False, keywords=False, corresponding=False
    )
    # No strong signal -> returns the original path unchanged, writes nothing.
    assert strip_front_matter_from_docx(path, tmp_path / "work") == path


def test_convert_body_identical_when_no_front_matter(tmp_path):
    """The strip flag must be a no-op for a document with no recognized title
    page: byte-identical body LaTeX with and without stripping."""
    path = _build_manuscript(
        tmp_path / "bare.docx", markers=False, abstract=False, keywords=False, corresponding=False
    )
    with_strip = convert_docx_to_body(path, tmp_path / "m1", strip_front_matter=True).tex
    without = convert_docx_to_body(path, tmp_path / "m2", strip_front_matter=False).tex
    assert with_strip == without


def test_convert_body_drops_title_page(tmp_path):
    path = _build_manuscript(tmp_path / "full.docx")
    body = convert_docx_to_body(path, tmp_path / "media", strip_front_matter=True).tex

    assert "Frustrated Coupling in Model Heterostructures" not in body
    assert "Alice Author" not in body
    assert "alice@example.edu" not in body
    assert "We report a study of frustrated" not in body
    # The real body content survives, including its section heading.
    assert "\\section{Introduction}" in body
    assert "magnetic frustration has a long history" in body


# --------------------------------------------------------------------------- #
# End-to-end through the emitter: title in metadata, NOT duplicated in body
# --------------------------------------------------------------------------- #


def _build_unstyled_manuscript(path: Path) -> Path:
    """A manuscript in the messy real-world shape the YIG paper exposed: NO
    Word title/heading styles, a large-font title, composite "1*" markers, an
    "ABSTRACT:" label (colon), and a bare ALL-CAPS "INTRODUCTION" section
    heading. Everything the clean-styled fixtures don't stress.
    """
    docx = pytest.importorskip("docx")
    from docx.shared import Pt

    doc = docx.Document()

    title = doc.add_paragraph()
    trun = title.add_run("UNDERSTANDING FRUSTRATED COUPLING IN HETEROSTRUCTURES")
    trun.font.size = Pt(18)  # largest font -> title (no Title style used)

    authors = doc.add_paragraph()
    authors.add_run("Alice Author")
    _add_superscript(authors, "1*")
    authors.add_run(", Bob Builder")
    _add_superscript(authors, "2")

    aff1 = doc.add_paragraph()
    _add_superscript(aff1, "1")
    aff1.add_run("Institute of Physics, Example University")

    aff2 = doc.add_paragraph()
    _add_superscript(aff2, "2")
    aff2.add_run("Department of Materials, Sample Institute")

    corr = doc.add_paragraph()
    _add_superscript(corr, "*")
    corr.add_run("To whom correspondence should be addressed. Email: alice@example.edu")

    doc.add_paragraph("ABSTRACT:")
    doc.add_paragraph(
        "We report a study of frustrated magnetic coupling in model "
        "heterostructures and its dependence on layer composition."
    )

    doc.add_paragraph("INTRODUCTION")  # bare all-caps section heading, no style
    doc.add_paragraph(
        "The study of magnetic frustration has a long history in condensed "
        "matter physics."
    )

    doc.save(path)
    return path


def test_unstyled_manuscript_strips_through_abstract(tmp_path):
    """The YIG-shaped case: unstyled title page + 'ABSTRACT:' + all-caps
    section heading. The whole title page (including the corresponding line and
    abstract) is stripped; the all-caps body heading and body survive."""
    path = _build_unstyled_manuscript(tmp_path / "unstyled.docx")

    span = front_matter_span(path)
    assert span is not None
    # p0 title, p1 authors, p2/p3 affiliations, p4 corresponding line,
    # p5 "ABSTRACT:", p6 abstract body -> body "INTRODUCTION" is p7.
    assert span == (0, 7)

    body = convert_docx_to_body(path, tmp_path / "media", strip_front_matter=True).tex
    assert "UNDERSTANDING FRUSTRATED COUPLING" not in body
    assert "alice@example.edu" not in body
    assert "We report a study of frustrated" not in body
    # The all-caps body heading and the body itself remain. (pandoc hard-wraps
    # lines, so match a fragment that can't straddle a wrap.)
    assert "INTRODUCTION" in body
    assert "long history in condensed matter" in body


def test_emit_title_in_metadata_not_body(tmp_path):
    from latextify.emit.project import emit_project

    path = _build_manuscript(tmp_path / "full.docx")
    result = emit_project(path, "revtex4-2", tmp_path / "output", report=False)

    body = result.body_tex_path.read_text(encoding="utf-8")
    metadata = result.metadata_tex_path.read_text(encoding="utf-8")

    # The journal template renders the title + authors from paper.yaml...
    assert "Frustrated Coupling in Model Heterostructures" in metadata
    assert "Alice Author" in metadata
    # ...and they no longer appear a SECOND time in the converted body.
    assert "Frustrated Coupling in Model Heterostructures" not in body
    assert "We report a study of frustrated" not in body
    # Body content is intact.
    assert "magnetic frustration has a long history" in body
