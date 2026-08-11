"""Figures the manuscript captions but has no image for (GUI_MULTIFILE item 5).

The bug this exists to surface is silent and ships wrong papers. Word's
caption convention binds a caption to the image beside it, and pandoc's docx
reader does the same -- so a caption-only "Figure 3" sitting above figure 4's
image gets bound to THAT image. VERIFIED on the fixtures below: a manuscript
captioning figures 1-4 with no image for 3 emits three figures, the third
being figure 4's image under figure 3's caption, with "Figure 4: ..." left in
the body as an ordinary paragraph. Before this detection existed, nothing
warned.

The number reported must be the genuinely missing one. A count-based check
("4 captions, 3 images -> the 4th is missing") names Figure 4 here, which
would send the author looking for the wrong file; these tests pin the
adjacency-based answer instead.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from latextify.emit.project import emit_project
from latextify.figures.caption_gaps import caption_gaps, stated_caption_numbers


def _png() -> io.BytesIO:
    buf = io.BytesIO()
    Image.new("RGB", (32, 24), (180, 60, 60)).save(buf, format="PNG")
    buf.seek(0)
    return buf


def _manuscript(path: Path, *, captions: int, missing: set[int] | None = None) -> Path:
    """A manuscript captioning 1..captions, omitting the image for `missing`."""
    missing = missing or set()
    document = Document()
    document.add_heading("Caption Gap Fixture", level=0)
    document.add_heading("Results", level=1)
    for number in range(1, captions + 1):
        document.add_paragraph(f"Result {number}.")
        if number not in missing:
            document.add_picture(_png(), width=Inches(1))
        document.add_paragraph(f"Figure {number}: Figure {number} text.", style="Caption")
    document.save(path)
    return path


@pytest.mark.parametrize("missing", [2, 3, 4])
def test_the_reported_gap_is_the_figure_actually_missing(tmp_path, missing):
    docx = _manuscript(tmp_path / f"gap{missing}.docx", captions=4, missing={missing})

    assert caption_gaps(docx, image_count=3) == [missing]


def test_a_manuscript_whose_captions_and_images_agree_is_silent(tmp_path):
    docx = _manuscript(tmp_path / "ok.docx", captions=3)

    assert caption_gaps(docx, image_count=3) == []


def test_two_missing_figures_are_both_reported(tmp_path):
    docx = _manuscript(tmp_path / "two.docx", captions=4, missing={2, 4})

    assert caption_gaps(docx, image_count=2) == [2, 4]


def test_non_contiguous_caption_numbering_is_left_alone(tmp_path):
    """An author numbering figures 2, 5, 9 is doing something deliberate; the
    heuristic must not flood them with false gaps."""
    document = Document()
    document.add_heading("Odd numbering", level=0)
    for number in (2, 5, 9):
        document.add_picture(_png(), width=Inches(1))
        document.add_paragraph(f"Figure {number}: text.", style="Caption")
    docx = tmp_path / "odd.docx"
    document.save(docx)

    assert caption_gaps(docx, image_count=3) == []


def test_a_manuscript_with_no_captions_is_silent(tmp_path):
    document = Document()
    document.add_heading("No captions", level=0)
    document.add_picture(_png(), width=Inches(1))
    docx = tmp_path / "nocaps.docx"
    document.save(docx)

    assert stated_caption_numbers(docx) == []
    assert caption_gaps(docx, image_count=1) == []


def test_an_unreadable_file_never_breaks_the_check(tmp_path):
    """Advisory detection must not be the thing that fails a conversion."""
    bogus = tmp_path / "not-a-docx.docx"
    bogus.write_bytes(b"certainly not a zip")

    assert stated_caption_numbers(bogus) == []
    assert caption_gaps(bogus, image_count=0) == []


def test_the_gap_surfaces_as_an_emit_warning_naming_the_remedy(tmp_path):
    docx = _manuscript(tmp_path / "gap.docx", captions=4, missing={3})

    result = emit_project(docx, "revtex4-2", tmp_path / "out")

    gaps = [w.message for w in result.warnings if "captions Figure" in w.message]
    assert len(gaps) == 1
    assert "Figure 3" in gaps[0]
    assert "figures/fig3." in gaps[0]  # names the file to supply
    assert "WRONG image" in gaps[0]  # says what actually goes wrong


def test_a_clean_manuscript_emits_no_caption_gap_warning(tmp_path):
    docx = _manuscript(tmp_path / "ok.docx", captions=2)

    result = emit_project(docx, "revtex4-2", tmp_path / "out")

    assert [w for w in result.warnings if "captions Figure" in w.message] == []
