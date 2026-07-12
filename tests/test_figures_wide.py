"""Figure width-bounding and single-vs-wide (two-column) float selection.

Real manuscripts embed figures at their native pixel size, and the emitter
used to write a bare ``\\includegraphics{...}`` with no width option -- so a
high-resolution figure rendered at natural size and ran clean off the page
(five "Float too large for page" warnings, up to 701pt of overflow, on the
first real manuscript). Every float figure must now be bounded to
``\\linewidth``. On top of that, landscape multi-panel composites (the norm in
physics papers) render unreadably small squeezed into one column of a
two-column journal, so a figure wider than :data:`_WIDE_ASPECT_THRESHOLD` is
emitted as the journal's wide float (``figure*``) to span both columns.

Figures are built with Pillow at run time -- no committed binary fixture,
nothing from any real manuscript.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from latextify.emit.project import _is_wide_figure, emit_project


def _png_bytes(size: tuple[int, int], color: tuple[int, int, int] = (40, 90, 160)) -> io.BytesIO:
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    buf.seek(0)
    return buf


def _docx_with_image(path, size: tuple[int, int]) -> None:
    document = Document()
    document.add_heading("Figure Aspect Fixture", level=0)
    document.add_heading("Results", level=1)
    document.add_paragraph("A placeholder result, illustrated below.")
    # Display width is fixed; only the source image's PIXEL aspect ratio should
    # drive the single-vs-wide decision, never the docx layout width.
    document.add_picture(_png_bytes(size), width=Inches(3))
    document.add_paragraph("Figure 1: A placeholder result.", style="Caption")
    document.save(path)


# --------------------------------------------------------------------------- #
# _is_wide_figure unit behavior (the classifier itself)
# --------------------------------------------------------------------------- #


def test_landscape_image_is_wide(tmp_path):
    p = tmp_path / "wide.png"
    Image.new("RGB", (400, 100), (10, 10, 10)).save(p)  # aspect 4.0
    assert _is_wide_figure(p) is True


@pytest.mark.parametrize("size", [(100, 400), (100, 100), (120, 100)])
def test_portrait_square_and_mild_landscape_are_not_wide(tmp_path, size):
    # 0.25 (portrait), 1.0 (square), 1.2 (mild landscape, below the 1.3 cut).
    p = tmp_path / "narrow.png"
    Image.new("RGB", size, (10, 10, 10)).save(p)
    assert _is_wide_figure(p) is False


def test_unreadable_file_degrades_to_not_wide(tmp_path):
    # Sizing must never fail a conversion that otherwise compiles: a file
    # Pillow cannot open (here a bogus "PDF") returns False, not an exception.
    p = tmp_path / "not-an-image.pdf"
    p.write_bytes(b"%PDF-1.4 not really an image\n")
    assert _is_wide_figure(p) is False


# --------------------------------------------------------------------------- #
# end-to-end: the chosen float environment + width bound in body.tex
# --------------------------------------------------------------------------- #


def test_every_float_figure_is_width_bounded(tmp_path):
    docx = tmp_path / "square.docx"
    _docx_with_image(docx, (64, 64))
    result = emit_project(docx, "revtex4-2", tmp_path / "output", report=False)
    body = result.body_tex_path.read_text(encoding="utf-8")

    # No bare \includegraphics survives -- overflow guard for the whole class.
    assert "\\includegraphics{" not in body
    assert "\\includegraphics[width=\\linewidth]{figures/fig1.png}" in body


def test_landscape_figure_spans_both_columns(tmp_path):
    docx = tmp_path / "wide.docx"
    _docx_with_image(docx, (400, 100))  # aspect 4.0 -> wide
    result = emit_project(docx, "revtex4-2", tmp_path / "output", report=False)
    body = result.body_tex_path.read_text(encoding="utf-8")

    assert "\\begin{figure*}" in body
    assert "\\end{figure*}" in body
    assert result.figures[0].wide is True


def test_portrait_figure_stays_single_column(tmp_path):
    docx = tmp_path / "tall.docx"
    _docx_with_image(docx, (100, 400))  # aspect 0.25 -> single column
    result = emit_project(docx, "revtex4-2", tmp_path / "output", report=False)
    body = result.body_tex_path.read_text(encoding="utf-8")

    assert "\\begin{figure}" in body
    assert "\\begin{figure*}" not in body
    assert result.figures[0].wide is False
