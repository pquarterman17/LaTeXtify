"""Generate tests/fixtures/tables.docx.

Exercises the table-normalization filter (plan item 17):

- A clean 3-column table with a header row and mixed numeric/text columns
  (a text ID column, a numeric temperature column, a text notes column) --
  the case that must convert to compiling booktabs LaTeX.
- A pathological table with a horizontal merge (``gridSpan``, header row)
  *and* a vertical merge (``vMerge``, body column) -- the case that must
  degrade to a :class:`~latextify.model.FilterFinding` warning and fall back
  to pandoc's own default table rendering rather than attempt (and risk
  silently corrupting) a reconstruction.

python-docx's ``_Cell.merge()`` authors both ``gridSpan`` (horizontal) and
``vMerge`` (vertical) OOXML directly -- no raw XML injection needed here,
unlike make_equations.py/make_unsupported.py.

Run with:
    uv run python tests/fixtures/make_tables.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURE_PATH = Path(__file__).parent / "tables.docx"


def _add_clean_table(doc: Document) -> None:
    doc.add_heading("Sample Measurements", level=1)
    doc.add_paragraph("The following samples were characterized at cryogenic temperature.")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "Sample"
    header[1].text = "Temperature (K)"
    header[2].text = "Notes"

    rows = [
        ("A", "4.2", "superconducting"),
        ("B", "77", "boiling N2"),
        ("C", "300", "room temperature"),
    ]
    for r, (sample, temp, notes) in enumerate(rows, start=1):
        cells = table.rows[r].cells
        cells[0].text = sample
        cells[1].text = temp
        cells[2].text = notes


def _add_pathological_table(doc: Document) -> None:
    doc.add_heading("Instrument Log", level=1)
    doc.add_paragraph(
        "This table has a merged header banner and a vertically merged "
        "cell; it must NOT be silently reconstructed."
    )

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Horizontal merge across columns 1-2 in the header row.
    table.cell(0, 0).text = "Run"
    merged_header = table.cell(0, 1).merge(table.cell(0, 2))
    merged_header.text = "Measurement"

    table.cell(1, 0).text = "1"
    table.cell(1, 2).text = "K"
    # Vertical merge across rows 1-2 in column 1 (vMerge).
    merged_body = table.cell(1, 1).merge(table.cell(2, 1))
    merged_body.text = "10"

    table.cell(2, 0).text = "2"
    table.cell(2, 2).text = "L"


def build() -> None:
    doc = Document()
    doc.add_heading("Table Normalization Fixture", level=1)

    _add_clean_table(doc)
    doc.add_paragraph("")
    _add_pathological_table(doc)

    doc.save(FIXTURE_PATH)
    print(f"wrote {FIXTURE_PATH}")


if __name__ == "__main__":
    build()
