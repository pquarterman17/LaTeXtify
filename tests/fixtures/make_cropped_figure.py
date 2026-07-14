"""Generate tests/fixtures/cropped_figure.docx.

One embedded image carrying a Word display crop (``a:srcRect``) so the crop
pipeline (``latextify.figures.extract`` reads it, ``latextify.figures.convert``
applies it) can be exercised end-to-end. Word crops an image for *display* but
keeps the full original pixels in ``word/media/``; without the fix those hidden
regions leak into the output.

The source image is a 100x100 PNG split into four solid colour quadrants:

    top-left  = red      top-right    = green
    bottom-left = blue    bottom-right = yellow

The injected ``a:srcRect`` hides the right and bottom halves (``r=50000``,
``b=50000`` -- thousandths of a percent), so the *visible* region is exactly the
top-left quadrant. A correct crop therefore yields a 50x50 all-red PNG; a
regression (no crop) leaves the full 100x100 four-colour image, which is trivial
to assert against.

Regenerate with::

    uv run python tests/fixtures/make_cropped_figure.py
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image

FIXTURE_PATH = Path(__file__).parent / "cropped_figure.docx"

RED = (220, 40, 40)
GREEN = (40, 160, 40)
BLUE = (40, 40, 220)
YELLOW = (220, 220, 40)


def _quadrant_png(size: int = 100) -> io.BytesIO:
    """A ``size``x``size`` PNG with four distinct solid-colour quadrants."""
    image = Image.new("RGB", (size, size))
    half = size // 2
    image.paste(RED, (0, 0, half, half))
    image.paste(GREEN, (half, 0, size, half))
    image.paste(BLUE, (0, half, half, size))
    image.paste(YELLOW, (half, half, size, size))
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _inject_srcrect(shape, *, left=0, top=0, right=50000, bottom=50000) -> None:
    """Insert an ``a:srcRect`` into the picture's ``pic:blipFill`` after its blip."""
    inline = shape._inline
    blip_fill = next(el for el in inline.iter() if el.tag == qn("pic:blipFill"))
    blip = blip_fill.find(qn("a:blip"))
    src_rect = OxmlElement("a:srcRect")
    for attr, value in (("l", left), ("t", top), ("r", right), ("b", bottom)):
        src_rect.set(attr, str(value))
    blip.addnext(src_rect)


def build() -> None:
    document = Document()
    document.add_heading("Cropped Figure Fixture", level=0)
    document.add_paragraph("The figure below is cropped in Word to its top-left quadrant.")
    shape = document.add_picture(_quadrant_png(), width=Inches(2))
    _inject_srcrect(shape)
    document.add_paragraph(
        "Figure 1: A quadrant image whose right and bottom halves are cropped away."
    )
    document.save(FIXTURE_PATH)
    print(f"wrote {FIXTURE_PATH}")


if __name__ == "__main__":
    build()
