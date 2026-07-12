"""Generate ``bracket_cited.docx`` -- a manuscript with a "[N]"-prefixed reference list.

Real papers frequently type their reference list as "[1] A. Author, Title,
Journal 12, 345 (2020)." rather than "1. A. Author, ..." (the shape
``hand_cited.docx`` exercises). This is the regression fixture for GAP 1: a
5-entry bracket-prefixed reference list, one entry (4) deliberately typed with
NO space after the closing bracket ("[4]Doe, ...") -- the exact real-world
shape that leaked a raw "[4]" into the reconstructed entry text before the
fix (observed key "4b2015" with title "{[}4]...").

In-text markers exercise the single/list/range forms:

* ``[2]``      -- single numeric marker
* ``[3,5]``    -- a numeric list
* ``[1-3]``    -- a numeric range

Every reference is deliberately matchable (a distinctive lowercase phrase
mocked to a full Crossref record in the test) -- GAP 1 is about the bracket
PREFIX parsing itself, not about reconciliation-confidence edge cases
(``hand_cited.docx`` already covers those).

Run directly to (re)write the fixture next to this script::

    uv run python tests/fixtures/make_bracket_cited.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURE_PATH = Path(__file__).with_name("bracket_cited.docx")

# (number, reference text as typed, WITH its bracket prefix baked in). Entry 4
# has no space after "]" -- the specific shape that used to leak the bracket.
REFERENCES: list[str] = [
    "[1] Alpha, A. Foundational widget calibration techniques. "
    "Journal of Widget Physics 10, 100-110 (2020).",
    "[2] Bravo, B. Superconducting widget arrays at low temperature. "
    "Journal of Widget Physics 11, 200-210 (2019).",
    "[3] Charlie, C. Magnetotransport in doped widget thin films. "
    "Journal of Widget Physics 12, 300-310 (2018).",
    "[4]Delta, D. Widget growth via molecular beam epitaxy. "
    "Journal of Widget Physics 13, 400-410 (2017).",
    "[5] Echo, E. Spectroscopic signatures of widget defects. "
    "Journal of Widget Physics 14, 500-510 (2016).",
]


def build(out_path: Path | str = FIXTURE_PATH) -> Path:
    """Write the .docx fixture and return its path."""
    out_path = Path(out_path)
    document = Document()
    document.add_heading("Bracket-Prefixed Reference List Reconstruction", level=0)  # Title

    document.add_heading("Introduction", level=1)
    document.add_paragraph(
        "The superconducting widget array platform [2] built on earlier "
        "calibration work. Subsequent studies [3,5] refined the growth "
        "process, and the foundational sequence of results [1-3] established "
        "the platform's core physics."
    )

    document.add_heading("References", level=1)
    for reference in REFERENCES:
        document.add_paragraph(reference)

    document.save(out_path)
    return out_path


if __name__ == "__main__":
    written = build()
    print(f"wrote {written}")
