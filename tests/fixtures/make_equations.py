"""Generate tests/fixtures/equations.docx.

Exercises the pandoc body pipeline (plan item 3): headings at levels 1-4
(level 4 to prove the >3 clamp in ``filters.normalize_headings`` fires) plus
inline and display OMML math (Word's equation-editor format).

python-docx cannot author OMML equations itself, so the ``<m:oMath>`` /
``<m:oMathPara>`` XML below is injected directly into paragraph runs via
python-docx's low-level ``parse_xml``/oxml API. The markup is not
hand-guessed: it is the exact OMML pandoc's own docx *writer* emits for
``$\\frac{a}{b}$`` / ``$$\\frac{a}{b}$$`` (verified by round-tripping a
markdown fixture through ``pandoc -t docx`` and inspecting
``word/document.xml``), i.e. schema-valid math identical in shape to what
Word's equation editor produces for a simple fraction.

Plan item 23 (equation audit tooling) extends this fixture with three more
real OMML shapes -- a matrix, a piecewise ``cases`` function, and a system of
aligned equations (``m:eqArr``) -- captured the same way (round-tripped
through pandoc's own docx writer for ``$$\\begin{pmatrix}...``,
``$$f(x) = \\begin{cases}...``, and ``$$\\begin{aligned}...`` respectively,
then lifted out of ``word/document.xml``), plus one deliberately exotic
construct: a literally empty ``<m:oMath/>`` -- a blank equation-editor
placeholder a real author left behind. All four are things pandoc's own
*writer* considers valid OMML for the equivalent LaTeX, so they convert back
cleanly; the empty one is the "interesting case" plan item 23 calls out --
it round-trips to an empty ``\\(\\)`` rather than being dropped, which is
exactly the kind of silent-content-loss risk the equation audit exists to
surface to a human.

Run with:
    uv run python tests/fixtures/make_equations.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

FIXTURE_PATH = Path(__file__).parent / "equations.docx"

# Real OMML for `a over b`, captured from pandoc's docx writer (see module
# docstring). `{jc}` is only present in the oMathPara (display) wrapper.
_OMATH_FRAC = (
    "<m:oMath>"
    '<m:f><m:fPr><m:type m:val="bar"/></m:fPr>'
    "<m:num><m:r><m:t>a</m:t></m:r></m:num>"
    "<m:den><m:r><m:t>b</m:t></m:r></m:den>"
    "</m:f>"
    "</m:oMath>"
)

# The three OMML fragments below are the *inner* content of an <m:oMath>
# element (i.e. with the outer <m:oMath>...</m:oMath> tags stripped, same
# convention as `_OMATH_FRAC` above once sliced in `build()`), captured by
# round-tripping through pandoc's own docx writer:
#   matrix : $$\begin{pmatrix} a & b \\ c & d \end{pmatrix}$$
#   cases  : $$f(x) = \begin{cases} 1 & x > 0 \\ 0 & x \le 0 \end{cases}$$
#   eqarray: $$\begin{aligned} a &= b + c \\ d &= e - f \end{aligned}$$
_OMATH_MATRIX_INNER = (
    '<m:d><m:dPr><m:begChr m:val="(" /><m:sepChr m:val="" /><m:endChr m:val=")" />'
    "<m:grow /></m:dPr><m:e><m:m><m:mPr><m:baseJc m:val=\"center\" />"
    '<m:plcHide m:val="on" /><m:mcs><m:mc><m:mcPr><m:mcJc m:val="center" />'
    '<m:count m:val="1" /></m:mcPr></m:mc><m:mc><m:mcPr><m:mcJc m:val="center" />'
    '<m:count m:val="1" /></m:mcPr></m:mc></m:mcs></m:mPr>'
    "<m:mr><m:e><m:r><m:t>a</m:t></m:r></m:e><m:e><m:r><m:t>b</m:t></m:r></m:e></m:mr>"
    "<m:mr><m:e><m:r><m:t>c</m:t></m:r></m:e><m:e><m:r><m:t>d</m:t></m:r></m:e></m:mr>"
    "</m:m></m:e></m:d>"
)

_OMATH_CASES_INNER = (
    '<m:r><m:t>f</m:t></m:r><m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>(</m:t></m:r>'
    '<m:r><m:t>x</m:t></m:r><m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>)</m:t></m:r>'
    '<m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>=</m:t></m:r>'
    '<m:d><m:dPr><m:begChr m:val="{" /><m:sepChr m:val="" /><m:endChr m:val="" />'
    '<m:grow /></m:dPr><m:e><m:m><m:mPr><m:baseJc m:val="center" />'
    '<m:plcHide m:val="on" /><m:mcs><m:mc><m:mcPr><m:mcJc m:val="left" />'
    '<m:count m:val="1" /></m:mcPr></m:mc><m:mc><m:mcPr><m:mcJc m:val="left" />'
    '<m:count m:val="1" /></m:mcPr></m:mc></m:mcs></m:mPr>'
    "<m:mr><m:e><m:r><m:t>1</m:t></m:r></m:e><m:e><m:r><m:t>x</m:t></m:r>"
    '<m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>&gt;</m:t></m:r>'
    "<m:r><m:t>0</m:t></m:r></m:e></m:mr>"
    "<m:mr><m:e><m:r><m:t>0</m:t></m:r></m:e><m:e><m:r><m:t>x</m:t></m:r>"
    '<m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>≤</m:t></m:r>'
    "<m:r><m:t>0</m:t></m:r></m:e></m:mr>"
    "</m:m></m:e></m:d>"
)

_OMATH_EQARRAY_INNER = (
    "<m:eqArr><m:e><m:r><m:t>a</m:t></m:r><m:r><m:t>&amp;</m:t></m:r>"
    '<m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>=</m:t></m:r>'
    '<m:r><m:t>b</m:t></m:r><m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>+</m:t></m:r>'
    "<m:r><m:t>c</m:t></m:r></m:e>"
    "<m:e><m:r><m:t>d</m:t></m:r><m:r><m:t>&amp;</m:t></m:r>"
    '<m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>=</m:t></m:r>'
    '<m:r><m:t>e</m:t></m:r><m:r><m:rPr><m:sty m:val="p" /></m:rPr><m:t>−</m:t></m:r>'
    "<m:r><m:t>f</m:t></m:r></m:e></m:eqArr>"
)


def _append_omath(paragraph, inner_xml: str) -> None:
    """Append an inline `<m:oMath>` equation as a direct child of the paragraph."""
    xml = f"<m:oMath {nsdecls('m')}>{inner_xml}</m:oMath>"
    paragraph._p.append(parse_xml(xml))


def _append_omath_para(paragraph, inner_xml: str) -> None:
    """Append a centered display `<m:oMathPara><m:oMath>...` block."""
    xml = (
        f"<m:oMathPara {nsdecls('m')}>"
        '<m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>'
        f"<m:oMath>{inner_xml}</m:oMath>"
        "</m:oMathPara>"
    )
    paragraph._p.append(parse_xml(xml))


def build() -> None:
    doc = Document()

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This manuscript exercises Word equation-editor math.")

    doc.add_heading("Method", level=2)
    p = doc.add_paragraph("The result follows from ")
    # Strip the OMML's own outer <m:oMath> wrapper tag by passing only the
    # inner fraction markup through _append_omath, which re-adds the wrapper.
    inner = _OMATH_FRAC[len("<m:oMath>") : -len("</m:oMath>")]
    _append_omath(p, inner)
    p.add_run(" as shown below.")

    doc.add_heading("Results", level=3)
    doc.add_paragraph("The final expression is:")
    display_p = doc.add_paragraph()
    _append_omath_para(display_p, inner)
    doc.add_paragraph("This concludes the derivation.")

    # Plan item 23: matrix, cases, and equation-array display math, plus one
    # deliberately exotic blank equation placeholder -- see module docstring.
    doc.add_heading("Special Forms", level=2)
    doc.add_paragraph("A two-by-two matrix:")
    matrix_p = doc.add_paragraph()
    _append_omath_para(matrix_p, _OMATH_MATRIX_INNER)

    doc.add_paragraph("A piecewise-defined function:")
    cases_p = doc.add_paragraph()
    _append_omath_para(cases_p, _OMATH_CASES_INNER)

    doc.add_paragraph("A system of aligned equations:")
    eqarray_p = doc.add_paragraph()
    _append_omath_para(eqarray_p, _OMATH_EQARRAY_INNER)

    blank_p = doc.add_paragraph("An equation the author started and abandoned: ")
    _append_omath(blank_p, "")  # deliberately exotic: a literally empty <m:oMath/>
    blank_p.add_run(" (nothing should appear between those parentheses).")

    # Heading 4 has no journal-template counterpart (classes stop at
    # \subsubsection); normalize_headings() must clamp it to level 3.
    doc.add_heading("Too Deep A Section", level=4)
    doc.add_paragraph("This paragraph lives under the clamped heading.")

    doc.save(FIXTURE_PATH)
    print(f"wrote {FIXTURE_PATH}")


if __name__ == "__main__":
    build()
