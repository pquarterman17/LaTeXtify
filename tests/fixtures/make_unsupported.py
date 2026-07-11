"""Generate tests/fixtures/unsupported.docx.

Plants exactly one construct per `latextify.ingest.preflight` detector:
text box, tracked-change insertion, floating (anchored) image, SmartArt
diagram, and an "Equation N:" paragraph carrying a pasted image instead of
OMML. Each construct lives in its own paragraph so the five detectors never
cross-trigger on the same element.

python-docx cannot author text boxes, tracked changes, or drawing anchors
directly, so this script builds the ordinary paragraphs with python-docx and
splices in the remaining constructs as raw OOXML via `docx.oxml.parse_xml`,
inserted directly into `document.element.body`.

Run with: uv run python tests/fixtures/make_unsupported.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

FIXTURE_PATH = Path(__file__).parent / "unsupported.docx"

_VML_NS = 'xmlns:v="urn:schemas-microsoft-com:vml"'

TEXT_BOX_XML = f"""
<w:p {nsdecls("w")} {_VML_NS}>
  <w:r>
    <w:pict>
      <v:shape>
        <v:textbox>
          <w:txbxContent>
            <w:p><w:r><w:t>Planted text box content.</w:t></w:r></w:p>
          </w:txbxContent>
        </v:textbox>
      </v:shape>
    </w:pict>
  </w:r>
</w:p>
""".strip()

TRACKED_CHANGE_XML = f"""
<w:p {nsdecls("w")}>
  <w:ins w:id="101" w:author="Reviewer" w:date="2026-01-01T00:00:00Z">
    <w:r><w:t>This sentence was inserted via tracked changes.</w:t></w:r>
  </w:ins>
</w:p>
""".strip()

FLOATING_OBJECT_XML = f"""
<w:p {nsdecls("w", "wp", "a")}>
  <w:r>
    <w:drawing>
      <wp:anchor behindDoc="0" distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" locked="0" layoutInCell="1" allowOverlap="1"
                 relativeHeight="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>
        <wp:extent cx="914400" cy="914400"/>
        <wp:wrapNone/>
        <wp:docPr id="201" name="FloatingObject"/>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"/>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>
""".strip()

SMARTART_XML = f"""
<w:p {nsdecls("w", "wp", "a")}>
  <w:r>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="914400" cy="914400"/>
        <wp:docPr id="202" name="SmartArtDiagram"/>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/diagram"/>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>
""".strip()

EQUATION_AS_IMAGE_XML = f"""
<w:p {nsdecls("w", "wp", "a")}>
  <w:r><w:t>Equation 1: pasted as a screenshot instead of OMML.</w:t></w:r>
  <w:r>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="914400" cy="457200"/>
        <wp:docPr id="203" name="EquationScreenshot"/>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"/>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>
""".strip()

PLANTED_CONSTRUCTS = (
    TEXT_BOX_XML,
    TRACKED_CHANGE_XML,
    FLOATING_OBJECT_XML,
    SMARTART_XML,
    EQUATION_AS_IMAGE_XML,
)


def build() -> None:
    document = Document()
    document.add_heading("Preflight Unsupported-Constructs Fixture", level=0)  # Title style
    document.add_heading("Planted Constructs", level=1)
    document.add_paragraph(
        "This fixture plants exactly one construct per preflight detector, below."
    )

    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))  # must stay the last child of w:body
    for raw_xml in PLANTED_CONSTRUCTS:
        element = parse_xml(raw_xml)
        if sect_pr is not None:
            sect_pr.addprevious(element)
        else:
            body.append(element)

    document.save(FIXTURE_PATH)
    print(f"wrote {FIXTURE_PATH}")


if __name__ == "__main__":
    build()
