"""Generate tests/fixtures/supplement.docx -- a small SI manuscript (plan item 21).

Built with python-docx for structure (headings, two captioned images, one
OMML equation -- the same techniques ``make_figures.py``/``make_equations.py``
use) plus hand-injected Zotero ``ADDIN ZOTERO_ITEM CSL_CITATION`` field-code
runs, appended directly onto a normally-built paragraph's ``_p`` element via
python-docx's low-level ``parse_xml``/oxml API (the same technique
``make_equations.py`` uses for OMML). ``citations.fields.extract_field_citations``
then recognizes them exactly like a real Zotero-cited manuscript --
``make_zotero_cited.py``'s hand-assembled-.docx approach is not reused
wholesale here because it does not embed images, and this fixture needs both.

Two citations, to exercise plan item 21's cross-document dedup:

    - one SHARED with ``zotero_cited.docx``'s first article citation (same
      DOI, ``10.1103/PhysRevB.101.045123``, different itemData otherwise --
      dedup keys off the DOI alone) -- proves ``merge_ref_entries`` collapses
      it to a single ``references.bib`` entry instead of writing a duplicate.
    - one NEW reference (a distinct DOI) unique to the SI -- proves a
      genuinely new SI reference is still added.

Run with:
    uv run python tests/fixtures/make_supplement.py
"""

from __future__ import annotations

import io
import json
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

FIXTURE_PATH = Path(__file__).parent / "supplement.docx"

# Same DOI as make_zotero_cited.py's ARTICLE payload -- the shared-by-DOI case.
SHARED_DOI = "10.1103/PhysRevB.101.045123"

SHARED_ARTICLE = {
    "citationID": "si-shared-1",
    "properties": {"formattedCitation": "[S1]", "plainCitation": "[S1]", "noteIndex": 0},
    "citationItems": [
        {
            "id": 201,
            "uris": ["http://zotero.org/users/1/items/SI0001"],
            "itemData": {
                "id": 201,
                "type": "article-journal",
                "title": "Quantum transport in GaAs heterostructures",
                "container-title": "Physical Review B",
                "DOI": SHARED_DOI,
                "volume": "101",
                "issue": "4",
                "page": "045123",
                "author": [
                    {"family": "Müller", "given": "Hans"},
                    {"family": "Nyström", "given": "Erik"},
                ],
                "issued": {"date-parts": [[2020, 1, 15]]},
            },
        }
    ],
}

NEW_ARTICLE = {
    "citationID": "si-new-2",
    "properties": {"formattedCitation": "[S2]", "plainCitation": "[S2]", "noteIndex": 0},
    "citationItems": [
        {
            "id": 202,
            "uris": ["http://zotero.org/users/1/items/SI0002"],
            "itemData": {
                "id": 202,
                "type": "article-journal",
                "title": "Extended magnetotransport analysis of doped heterostructures",
                "container-title": "Physical Review Applied",
                "DOI": "10.1103/PhysRevApplied.15.054001",
                "volume": "15",
                "page": "054001",
                "author": [{"family": "Okafor", "given": "Chidi"}],
                "issued": {"date-parts": [[2021]]},
            },
        }
    ],
}


# --- tiny solid-color PNG (same approach as make_figures.py) ----------------


def _solid_png(color: tuple[int, int, int], size: int = 4) -> io.BytesIO:
    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0)  # 8-bit RGB
    raw_scanlines = b"".join(b"\x00" + bytes(color) * size for _ in range(size))
    idat = zlib.compress(raw_scanlines)
    png = signature + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")
    return io.BytesIO(png)


# --- OMML equation (same fraction shape as make_equations.py) ---------------

_OMATH_FRAC_INNER = (
    '<m:f><m:fPr><m:type m:val="bar"/></m:fPr>'
    "<m:num><m:r><m:t>dS</m:t></m:r></m:num>"
    "<m:den><m:r><m:t>dT</m:t></m:r></m:den>"
    "</m:f>"
)


def _append_omath(paragraph, inner_xml: str) -> None:
    xml = f"<m:oMath {nsdecls('m')}>{inner_xml}</m:oMath>"
    paragraph._p.append(parse_xml(xml))


# --- Zotero complex-field citation, appended onto a python-docx paragraph ---


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _fldchar_xml(char_type: str) -> str:
    return f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="{char_type}"/></w:r>'


def _instr_run_xml(text: str) -> str:
    return (
        f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve">'
        f"{_xml_escape(text)}</w:instrText></w:r>"
    )


def _text_run_xml(text: str) -> str:
    return f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_xml_escape(text)}</w:t></w:r>'


def _zotero_instruction(payload: dict) -> str:
    return " ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(payload, ensure_ascii=False) + " "


def _append_zotero_field(paragraph, payload: dict, result_text: str) -> None:
    """Append a complex Zotero citation field's runs onto ``paragraph``."""
    paragraph._p.append(parse_xml(_fldchar_xml("begin")))
    paragraph._p.append(parse_xml(_instr_run_xml(_zotero_instruction(payload))))
    paragraph._p.append(parse_xml(_fldchar_xml("separate")))
    paragraph._p.append(parse_xml(_text_run_xml(result_text)))
    paragraph._p.append(parse_xml(_fldchar_xml("end")))


def build() -> None:
    doc = Document()

    doc.add_heading("Supplementary Figures", level=1)
    doc.add_paragraph("An additional measurement is shown below.")
    doc.add_picture(_solid_png((220, 200, 40)), width=Inches(1))
    doc.add_paragraph(
        "Figure 1: An additional yellow-highlighted result not shown in the main text.",
        style="Caption",
    )

    doc.add_paragraph("A second supplementary measurement follows.")
    doc.add_picture(_solid_png((150, 40, 200)), width=Inches(1))
    doc.add_paragraph(
        "Figure 2: A purple placeholder figure, captioned via a plain paragraph."
    )

    doc.add_heading("Supplementary Equation", level=1)
    p = doc.add_paragraph("The temperature dependence of the entropy is given by ")
    _append_omath(p, _OMATH_FRAC_INNER)
    p.add_run(" as derived from the main text's free-energy expression.")

    doc.add_heading("Supplementary Discussion", level=1)
    p1 = doc.add_paragraph("As previously reported ")
    _append_zotero_field(p1, SHARED_ARTICLE, "[S1]")
    p1.add_run(", the transport in this material family is well characterized.")

    p2 = doc.add_paragraph("A complementary high-field measurement ")
    _append_zotero_field(p2, NEW_ARTICLE, "[S2]")
    p2.add_run(" supports the conclusions drawn in the main text.")

    doc.save(FIXTURE_PATH)
    print(f"wrote {FIXTURE_PATH}")


if __name__ == "__main__":
    build()
