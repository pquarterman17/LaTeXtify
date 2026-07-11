"""Generate tests/fixtures/clean.docx.

A well-behaved manuscript stub with no unsupported constructs: Title,
Heading 1/2, a Caption-styled paragraph, and plain body text only. Exercises
the "clean fixture yields zero errors" side of the preflight done-when
criteria, and gives the style-inventory detector real positive signal
(heading levels {1, 2}, Title used, Caption used) to check against.

Run with: uv run python tests/fixtures/make_clean.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURE_PATH = Path(__file__).parent / "clean.docx"


def build() -> None:
    document = Document()
    document.add_heading("A Perfectly Ordinary Manuscript", level=0)  # Title style
    document.add_heading("Introduction", level=1)
    document.add_paragraph(
        "This manuscript uses only plain paragraphs and standard Word styles, "
        "with nothing for the preflight detectors to flag."
    )
    document.add_heading("Background", level=2)
    document.add_paragraph("More ordinary body text, spanning a second paragraph.")
    document.add_paragraph("Figure 1: A perfectly ordinary caption.", style="Caption")

    document.save(FIXTURE_PATH)
    print(f"wrote {FIXTURE_PATH}")


if __name__ == "__main__":
    build()
