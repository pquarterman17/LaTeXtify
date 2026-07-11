"""Generate tests/fixtures/figures.docx.

Three embedded images, each captioned a different way so
``latextify.figures.extract`` exercises both association paths documented in
plan item 9's context (see the item 3 finding recorded there):

    Figure 1 -- Word's built-in "Caption" style, the common case. pandoc
        promotes this into a native ``Figure`` AST block and (in current
        pandoc versions) correctly carries the caption text through --
        exercises the direct ``Figure.caption`` path.
    Figure 2 -- a plain paragraph reading "Figure 2: ..." with no special
        style. pandoc leaves the image as a bare ``Image`` with no ``Figure``
        promotion, so the caption must be recovered from the adjacent
        sibling paragraph -- exercises the regex-fallback path that mitigates
        the item 3 finding (a Figure block with an empty pandoc-derived
        caption falls back the same way).
    Figure 3 -- a plain paragraph using the abbreviated "Fig. 3: ..." label,
        to exercise the alternate label spelling the caption regex accepts.

Source images are tiny solid-color PNGs synthesized in-script (no binary
asset files to commit) via raw PNG chunk encoding -- python-docx's picture
API accepts a file-like object, so nothing is written to disk first.

Regenerate with::

    uv run python tests/fixtures/make_figures.py
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches

FIXTURE_PATH = Path(__file__).parent / "figures.docx"


def _solid_png(color: tuple[int, int, int], size: int = 4) -> io.BytesIO:
    """A tiny solid-color RGB PNG, built by hand so no image library is needed."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0)  # 8-bit RGB
    raw_scanlines = b"".join(b"\x00" + bytes(color) * size for _ in range(size))
    idat = zlib.compress(raw_scanlines)
    png = signature + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")
    return io.BytesIO(png)


def build() -> None:
    document = Document()
    document.add_heading("Figures Fixture", level=0)  # Title style
    document.add_heading("Results", level=1)

    document.add_paragraph("The first result is shown below.")
    document.add_picture(_solid_png((220, 40, 40)), width=Inches(1))
    document.add_paragraph(
        "Figure 1: A red placeholder figure, captioned via Word's Caption style.",
        style="Caption",
    )

    document.add_paragraph("The second result follows.")
    document.add_picture(_solid_png((40, 160, 40)), width=Inches(1))
    document.add_paragraph(
        "Figure 2: A green placeholder figure, captioned via a plain paragraph."
    )

    document.add_paragraph("The third and final result is shown last.")
    document.add_picture(_solid_png((40, 40, 220)), width=Inches(1))
    document.add_paragraph(
        "Fig. 3: A blue placeholder figure, captioned with the abbreviated label."
    )

    document.save(FIXTURE_PATH)
    print(f"wrote {FIXTURE_PATH}")


if __name__ == "__main__":
    build()
