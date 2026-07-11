"""Generate tests/fixtures/metadata_titlepage.docx.

Regenerate with::

    uv run python tests/fixtures/make_metadata_titlepage.py

Requires python-docx (dev dependency only -- production code parses docx XML
directly with lxml; see latextify/ingest/metadata_guess.py).

Layout exercised (mirrors plan item 8's heuristics):
    1. Title-styled paragraph.
    2. Author line: two names, with superscript affiliation markers ("1",
       "2") and a superscript "*" marking the corresponding author.
    3. Two affiliation paragraphs, each prefixed by its numeric marker.
    4. A "Corresponding author: ..." line (marked "*") with an email.
    5. An "Abstract" paragraph (the heading text itself) followed by the
       abstract body.
    6. A "Keywords:" line.
"""

from pathlib import Path

from docx import Document

OUTPUT = Path(__file__).with_name("metadata_titlepage.docx")


def _add_superscript(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.superscript = True


def build() -> None:
    doc = Document()

    title = doc.add_paragraph(style="Title")
    title.add_run("Superconducting Gap Anisotropy in Doped Compound X2Y")

    authors = doc.add_paragraph()
    authors.add_run("Jane A. Doe")
    _add_superscript(authors, "1,*")
    authors.add_run(", John B. Smith")
    _add_superscript(authors, "1,2")

    aff1 = doc.add_paragraph()
    _add_superscript(aff1, "1")
    aff1.add_run("Department of Physics, University X, Springfield, USA")

    aff2 = doc.add_paragraph()
    _add_superscript(aff2, "2")
    aff2.add_run("Institute of Materials Science, Example City, USA")

    corr = doc.add_paragraph()
    _add_superscript(corr, "*")
    corr.add_run("Corresponding author: jane.doe@example.edu")

    doc.add_paragraph("Abstract")
    doc.add_paragraph(
        "We report magnetometry measurements of the doped compound X2Y "
        "revealing an anisotropic superconducting gap consistent with "
        "unconventional pairing symmetry."
    )

    doc.add_paragraph(
        "Keywords: superconductivity, magnetometry, doped compounds, gap anisotropy"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build()
