"""Generate ``hand_cited.docx`` -- a manuscript with typed (non-field) citations.

This is the plain-text reconstruction fixture (plan item 14): NO Zotero/Mendeley
field codes at all, just literal in-text markers and a typed, numbered reference
list after a "References" heading. It exercises every marker form the linker must
recognize in pandoc's LaTeX output:

* ``[1]`` and ``[12]``            -- single numeric markers
* ``[3-5,8]``                     -- a numeric range + list
* a superscript ``2,4`` run       -- superscript-numeral markers
* ``(Smith et al., 2020)``        -- an author-year marker (resolves to ref 1)

The reference list has 12 numbered entries. Entry 7 is deliberately obscure
("unpublished notes") so a Crossref lookup returns nothing and it is flagged for
verification -- the other 11 (>= 80%) reconstruct with DOIs under the mocked
Crossref in the tests.

Run directly to (re)write the fixture next to this script::

    uv run python tests/fixtures/make_hand_cited.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURE_PATH = Path(__file__).with_name("hand_cited.docx")

# (number, reference text). Entry 7 is the intentionally unmatchable one.
REFERENCES: list[str] = [
    "Smith, A. B., Jones, C. D. Coherent control of solid-state spin qubits. "
    "Nature Physics 16, 1201-1210 (2020).",
    "Anderson, P. W. Observation of topological superconductivity in a planar "
    "Josephson junction. Physical Review Letters 115, 020501 (2015).",
    "Brown, R. J., Lee, S. Magnetotransport signatures of Weyl semimetals. "
    "Physical Review B 98, 035001 (2018).",
    "Chen, X., Wang, Y. Two-dimensional magnets and their van der Waals "
    "heterostructures. Nature Materials 18, 1298-1310 (2019).",
    "Davis, M. K. Quantum oscillations in underdoped high-Tc cuprates. "
    "Science 356, 512-516 (2017).",
    "Evans, L., Patel, N. Moire flat bands in magic-angle twisted bilayer "
    "graphene. Reviews of Modern Physics 93, 025006 (2021).",
    "Foster, G. Unpublished laboratory notes on cryogenic amplifier noise, "
    "internal memorandum, 2016.",
    "Garcia, H., Ito, K. Spin-orbit torque switching of perpendicular "
    "magnetization. Nature Nanotechnology 9, 548-554 (2014).",
    "Hughes, T. L. Majorana zero modes in semiconductor nanowires. "
    "Science 339, 1057-1060 (2013).",
    "Ito, K., Nakamura, S. Room-temperature superconductivity in a hydride "
    "under high pressure. Nature 601, 35-40 (2022).",
    "Johnson, D. R. Berry-phase effects on electronic transport properties. "
    "Reviews of Modern Physics 84, 1419-1475 (2012).",
    "Novak, V., Petrov, A. Anomalous Hall effect in itinerant ferromagnets. "
    "Reviews of Modern Physics 83, 1539-1592 (2011).",
]


def build(out_path: Path | str = FIXTURE_PATH) -> Path:
    """Write the .docx fixture and return its path."""
    out_path = Path(out_path)
    document = Document()
    document.add_heading("Reconstructing Citations from Plain Text", level=0)  # Title

    document.add_heading("Introduction", level=1)
    document.add_paragraph(
        "The coherent control of spin qubits was first demonstrated [1] and later "
        "extended to larger arrays [12]. Several groups [3-5,8] have explored related "
        "materials platforms."
    )

    p = document.add_paragraph("Recent reviews")
    sup = p.add_run("2,4")
    sup.font.superscript = True
    p.add_run(" summarize the experimental landscape in detail.")

    document.add_paragraph(
        "The coherent-control protocol (Smith et al., 2020) underpins much of the "
        "subsequent work discussed here."
    )

    document.add_heading("References", level=1)
    for index, reference in enumerate(REFERENCES, start=1):
        document.add_paragraph(f"{index}. {reference}")

    document.save(out_path)
    return out_path


if __name__ == "__main__":
    written = build()
    print(f"wrote {written}")
