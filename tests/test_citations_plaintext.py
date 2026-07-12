"""Plain-text marker detection, reference-list segmentation, and body linkage.

Pure-Python tests that need no Crossref network access: the linkage tests build a
:class:`PlaintextResult` directly, and the segmentation tests build tiny .docx
files with python-docx.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from latextify.citations.plaintext import (
    PlaintextResult,
    _raw_leading_surname,
    expand_numeric_range,
    link_body_markers,
    segment_reference_list,
    strip_reference_section,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_CONTENT_TYPES = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    '<Default Extension="rels" '
    'ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
    '<Default Extension="xml" ContentType="application/xml"/>'
    '<Override PartName="/word/document.xml" '
    'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
    "</Types>"
)
_ROOT_RELS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rId1" '
    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
    'Target="word/document.xml"/>'
    "</Relationships>"
)


def _numbered_paragraph(text: str) -> str:
    """A paragraph carrying real Word list numbering (``w:pPr/w:numPr``).

    This is what Word's "Numbering" toolbar button produces: the displayed
    "1." is rendered by Word from the list definition, never typed as text --
    unlike ``_make_docx``'s ``f"{i}. {ref}"`` convenience below, which types
    the digits literally and does not exercise this path.
    """
    return (
        "<w:p><w:pPr><w:numPr><w:ilvl w:val=\"0\"/><w:numId w:val=\"1\"/></w:numPr></w:pPr>"
        f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    )


def _plain_paragraph(text: str) -> str:
    return f'<w:p><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'


def _heading_paragraph(text: str) -> str:
    return (
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
        f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    )


def _build_raw_docx(path: Path, body_paragraphs: list[str]) -> Path:
    body = "".join(body_paragraphs) + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/></w:sectPr>'
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:document xmlns:w="{W}"><w:body>{body}</w:body></w:document>'
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _CONTENT_TYPES)
        archive.writestr("_rels/.rels", _ROOT_RELS)
        archive.writestr("word/document.xml", document_xml)
    return path

# --------------------------------------------------------------------------- #
# numeric range expansion
# --------------------------------------------------------------------------- #


def test_expand_single():
    assert expand_numeric_range("12") == [12]


def test_expand_range_and_list():
    assert expand_numeric_range("3-5,8") == [3, 4, 5, 8]


def test_expand_multiple_ranges():
    assert expand_numeric_range("1-3, 7, 9-10") == [1, 2, 3, 7, 9, 10]


def test_expand_unicode_dash():
    assert expand_numeric_range("3–5") == [3, 4, 5]  # en dash


def test_expand_non_numeric_returns_empty():
    assert expand_numeric_range("see note") == []


def test_expand_reversed_range_kept_as_endpoints():
    assert expand_numeric_range("5-3") == [5, 3]


def test_expand_leading_zero_number_is_excluded():
    # "001" is a Miller/crystallographic index, e.g. "grown along [001]" --
    # never a citation; a typed reference list is never zero-padded.
    assert expand_numeric_range("001") == []


def test_expand_leading_zero_range_is_excluded():
    assert expand_numeric_range("001-005") == []


def test_expand_mixed_leading_zero_and_plain_number():
    # Only the leading-zero chunk is dropped; a genuine list entry survives.
    assert expand_numeric_range("1,001") == [1]


def test_expand_miller_index_triads_excluded():
    # "110" and "111" have no leading zero but are the remaining canonical
    # low-index cubic crystallographic directions -- see _MILLER_INDEX_TRIADS.
    assert expand_numeric_range("110") == []
    assert expand_numeric_range("111") == []
    assert expand_numeric_range("101") == []
    assert expand_numeric_range("011") == []


def test_expand_two_digit_numbers_starting_with_one_are_unaffected():
    # A genuine 2-digit reference number is NOT swept up by the Miller-index
    # exclusion -- that set is scoped to exactly the 3-digit triads.
    assert expand_numeric_range("10") == [10]
    assert expand_numeric_range("11") == [11]


# --------------------------------------------------------------------------- #
# reference-list segmentation
# --------------------------------------------------------------------------- #


def _make_docx(path: Path, heading: str, refs: list[str], *, numbered=True) -> Path:
    doc = Document()
    doc.add_heading("A Title", level=0)
    doc.add_heading("Body", level=1)
    doc.add_paragraph("Some body text citing [1] and [2].")
    doc.add_heading(heading, level=1)
    for i, ref in enumerate(refs, start=1):
        doc.add_paragraph(f"{i}. {ref}" if numbered else ref)
    doc.save(path)
    return path


def test_segment_finds_numbered_references(tmp_path):
    docx = _make_docx(
        tmp_path / "r.docx",
        "References",
        ["Smith, A. First. J. A 1 (2020).", "Jones, B. Second. J. B 2 (2019)."],
    )
    reflist = segment_reference_list(docx)
    assert reflist.found
    assert reflist.heading == "References"
    assert [r.number for r in reflist.references] == [1, 2]
    assert reflist.references[0].text.startswith("Smith")
    assert "1." not in reflist.references[0].text  # leading number stripped


def test_segment_bibliography_heading(tmp_path):
    docx = _make_docx(tmp_path / "b.docx", "Bibliography", ["Doe, J. Only. J. 1 (2021)."])
    reflist = segment_reference_list(docx)
    assert reflist.found
    assert reflist.heading == "Bibliography"


def test_segment_unnumbered_references(tmp_path):
    docx = _make_docx(
        tmp_path / "u.docx",
        "References",
        ["Smith, A. First paper. (2020).", "Jones, B. Second paper. (2019)."],
        numbered=False,
    )
    reflist = segment_reference_list(docx)
    assert reflist.found
    assert [r.number for r in reflist.references] == [None, None]


def test_segment_bracketed_numbers(tmp_path):
    doc = Document()
    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] Smith, A. First. (2020).")
    doc.add_paragraph("[2] Jones, B. Second. (2019).")
    path = tmp_path / "brk.docx"
    doc.save(path)
    reflist = segment_reference_list(path)
    assert [r.number for r in reflist.references] == [1, 2]
    assert reflist.references[0].text.startswith("Smith")


def test_segment_bracketed_numbers_no_space_after_bracket(tmp_path):
    # Real manuscripts often type "[4]Giles, ..." with NO space after the
    # closing bracket. Without tolerating this, the whole paragraph fails to
    # match at all: ref_number stays None and the raw "[4]" leaks into the
    # text handed to Crossref/raw-entry emission (observed real-world bug).
    doc = Document()
    doc.add_heading("References", level=1)
    doc.add_paragraph("[4]Giles, B. L. Fourth. (2020).")
    path = tmp_path / "brk_nospace.docx"
    doc.save(path)
    reflist = segment_reference_list(path)
    assert [r.number for r in reflist.references] == [4]
    assert reflist.references[0].text == "Giles, B. L. Fourth. (2020)."
    assert "[4]" not in reflist.references[0].text


def test_segment_parenthesized_numbers(tmp_path):
    doc = Document()
    doc.add_heading("References", level=1)
    doc.add_paragraph("(1) Smith, A. First. (2020).")
    doc.add_paragraph("(2) Jones, B. Second. (2019).")
    path = tmp_path / "paren.docx"
    doc.save(path)
    reflist = segment_reference_list(path)
    assert [r.number for r in reflist.references] == [1, 2]
    assert reflist.references[0].text.startswith("Smith")
    assert "(1)" not in reflist.references[0].text


def test_segment_parenthesized_numbers_no_space(tmp_path):
    doc = Document()
    doc.add_heading("References", level=1)
    doc.add_paragraph("(3)Doe, C. Third. (2018).")
    path = tmp_path / "paren_nospace.docx"
    doc.save(path)
    reflist = segment_reference_list(path)
    assert [r.number for r in reflist.references] == [3]
    assert reflist.references[0].text == "Doe, C. Third. (2018)."


def test_segment_decimal_at_paragraph_start_not_misread_as_reference_number(tmp_path):
    # Regression guard for the bracket/paren whitespace relaxation: a bare
    # "N." form must still require a trailing space, so "3.14 times as
    # large" at a paragraph's start is never misread as reference number 3.
    doc = Document()
    doc.add_heading("References", level=1)
    doc.add_paragraph("3.14159 is not a reference number. (2020).")
    path = tmp_path / "decimal.docx"
    doc.save(path)
    reflist = segment_reference_list(path)
    assert reflist.references[0].number is None
    assert reflist.references[0].text.startswith("3.14159")


def test_segment_no_reference_list(tmp_path):
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Body text with no reference section at all.")
    path = tmp_path / "none.docx"
    doc.save(path)
    reflist = segment_reference_list(path)
    assert not reflist.found
    assert reflist.heading is None
    assert reflist.references == []


def test_segment_skips_empty_paragraphs(tmp_path):
    doc = Document()
    doc.add_heading("References", level=1)
    doc.add_paragraph("1. Smith, A. First. (2020).")
    doc.add_paragraph("")  # blank line between references
    doc.add_paragraph("2. Jones, B. Second. (2019).")
    path = tmp_path / "gap.docx"
    doc.save(path)
    reflist = segment_reference_list(path)
    assert [r.number for r in reflist.references] == [1, 2]


def test_segment_word_native_numbered_list(tmp_path):
    # Real Word list numbering (w:numPr) -- the toolbar "Numbering" button --
    # displays "1.", "2.", ... without ever putting that text in a w:t run.
    # Without recognizing w:numPr, every reference gets number=None and every
    # in-text numeric marker in the body fails to link (keys_by_number stays
    # empty), which is a much more common real-world case than typed digits.
    docx = _build_raw_docx(
        tmp_path / "numpr.docx",
        [
            _heading_paragraph("References"),
            _numbered_paragraph("Smith, A. First paper. J. A 1 (2020)."),
            _numbered_paragraph("Jones, B. Second paper. J. B 2 (2019)."),
            _numbered_paragraph("Doe, C. Third paper. J. C 3 (2018)."),
        ],
    )
    reflist = segment_reference_list(docx)
    assert reflist.found
    assert [r.number for r in reflist.references] == [1, 2, 3]
    assert reflist.references[0].text.startswith("Smith")
    assert reflist.references[1].text.startswith("Jones")


def test_segment_word_native_numbered_list_mixed_with_typed_number(tmp_path):
    # A typed leading number always wins over the numPr-based sequential
    # fallback, even inside an otherwise auto-numbered list.
    docx = _build_raw_docx(
        tmp_path / "mixed.docx",
        [
            _heading_paragraph("References"),
            _numbered_paragraph("Smith, A. First paper. (2020)."),
            "<w:p><w:pPr><w:numPr><w:ilvl w:val=\"0\"/><w:numId w:val=\"1\"/></w:numPr></w:pPr>"
            '<w:r><w:t xml:space="preserve">99. Jones, B. Second paper. (2019).</w:t></w:r></w:p>',
        ],
    )
    reflist = segment_reference_list(docx)
    assert [r.number for r in reflist.references] == [1, 99]
    assert reflist.references[1].text.startswith("Jones")


def test_segment_plain_paragraphs_without_numpr_still_unnumbered(tmp_path):
    # No regression: ordinary (non-list, non-typed-number) paragraphs keep
    # number=None, same as before w:numPr recognition was added.
    docx = _build_raw_docx(
        tmp_path / "plain.docx",
        [
            _heading_paragraph("References"),
            _plain_paragraph("Smith, A. First paper. (2020)."),
            _plain_paragraph("Jones, B. Second paper. (2019)."),
        ],
    )
    reflist = segment_reference_list(docx)
    assert [r.number for r in reflist.references] == [None, None]


# --------------------------------------------------------------------------- #
# body linkage
# --------------------------------------------------------------------------- #


def _result(**over) -> PlaintextResult:
    defaults = dict(
        keys_by_number={
            1: "smith2020",
            2: "jones2019",
            3: "brown2018",
            4: "chen2017",
            8: "lee2015",
        },
        author_year_keys={("smith", "2020"): ["smith2020"]},
        has_reference_list=True,
        heading="References",
    )
    defaults.update(over)
    return PlaintextResult(**defaults)


def test_link_single_numeric_marker():
    tex, warnings = link_body_markers("Shown {[}1{]} here.", _result())
    assert tex == "Shown \\cite{smith2020} here."
    assert warnings == []


def test_link_numeric_range_marker():
    tex, warnings = link_body_markers("Groups {[}3-5,8{]} explored.", _result())
    # 5 is absent from keys_by_number -> partial, but 3,4,8 still link.
    assert "\\cite{brown2018,chen2017,lee2015}" in tex
    assert any("no reference numbered 5" in w for w in warnings)


def test_link_numeric_range_within_one_bracket_double_dash():
    # A single typed marker "[1-3]" with an en dash reaches this stage as the
    # body "{[}1--3{]}" (pandoc renders en dash -> "--"). The range's endpoints
    # must still expand -- previously "1--3" split to ["1", "", "3"] and the
    # whole marker was dropped, leaving literal "[1-3]" (the real-manuscript
    # "[8-10]"/"[11-13]"/"[19-23]" bug).
    tex, warnings = link_body_markers("Shown {[}1--3{]} here.", _result())
    assert "Shown \\cite{smith2020,jones2019,brown2018} here." == tex
    assert warnings == []


def test_link_numeric_range_within_one_bracket_em_dash():
    # Same, with an em dash pandoc renders as "---".
    tex, warnings = link_body_markers("Shown {[}2---4{]} here.", _result())
    assert "\\cite{jones2019,brown2018,chen2017}" in tex
    assert warnings == []


def test_link_numeric_range_across_separate_brackets():
    # pandoc brace-protects EACH bracket individually, so a typed "[1]-[3]"
    # range renders as two separate {[}N{]} groups joined by "--", not one
    # {[}1-3{]} group. Without merging them first, refs 2 (the range's
    # middle) is silently dropped -- confirmed against real pandoc 3.9 output
    # for "[1]–[3]" -> "{[}1{]}--{[}3{]}".
    tex, warnings = link_body_markers("See refs {[}1{]}--{[}3{]} for details.", _result())
    assert "\\cite{smith2020,jones2019,brown2018}" in tex
    assert warnings == []


def test_link_numeric_range_across_separate_brackets_unicode_endash():
    # Same case with a literal (unescaped) unicode en dash between brackets.
    tex, warnings = link_body_markers("See refs {[}1{]}–{[}3{]} for details.", _result())
    assert "\\cite{smith2020,jones2019,brown2018}" in tex
    assert warnings == []


def test_link_endnote_temporary_citation_resolves_via_author_year():
    # "{Davies, 2004 #78}" reaches this stage pandoc-escaped as
    # "\{Davies, 2004 \#78\}"; resolve it through the author-year index.
    result = _result(author_year_keys={("davies", "2004"): ["j2004"]})
    tex, warnings = link_body_markers("measurements\\{Davies, 2004 \\#78\\}.", result)
    assert tex == "measurements\\cite{j2004}."
    assert warnings == []


def test_link_endnote_temporary_citation_run_collapses_to_one_cite():
    # A tripled paste of the same unformatted field must become ONE \cite.
    result = _result(author_year_keys={("davies", "2004"): ["j2004"]})
    body = (
        "measurements\\{Davies, 2004 \\#78\\}"
        "\\{Davies, 2004 \\#78\\}\\{Davies, 2004 \\#78\\}."
    )
    tex, warnings = link_body_markers(body, result)
    assert tex == "measurements\\cite{j2004}."


def test_link_endnote_temporary_citation_unresolved_warns_and_stays_literal():
    tex, warnings = link_body_markers("x\\{Nobody, 1999 \\#5\\}.", _result())
    assert "\\{Nobody, 1999 \\#5\\}" in tex  # never fabricated into a \cite
    assert "\\cite" not in tex
    assert any("EndNote temporary citation" in w for w in warnings)


def test_raw_leading_surname_parses_typed_author():
    assert _raw_leading_surname("J. E. Davies, O. Hellwig, and K. Liu, ... (2004).") == "davies"
    assert _raw_leading_surname("B. J. Kirby, P. A. Kienzle, ...") == "kirby"
    assert _raw_leading_surname("") is None


def test_link_superscript_range_across_separate_commands():
    # Same splitting hazard for superscript markers: pandoc renders
    # "text^1^--^3^" as two separate \textsuperscript commands.
    tex, warnings = link_body_markers(
        "Reviews\\textsuperscript{1}--\\textsuperscript{3} summarize.", _result()
    )
    assert "\\cite{smith2020,jones2019,brown2018}" in tex
    assert warnings == []


def test_link_superscript_marker():
    tex, warnings = link_body_markers("Reviews\\textsuperscript{1,2} summarize.", _result())
    assert tex == "Reviews\\cite{smith2020,jones2019} summarize."
    assert warnings == []


def test_link_author_year_marker():
    tex, warnings = link_body_markers("The protocol (Smith et al., 2020) is key.", _result())
    assert tex == "The protocol \\cite{smith2020} is key."
    assert warnings == []


def test_link_author_year_across_line_wrap():
    # pandoc wraps long lines; the marker may straddle a newline.
    tex, _ = link_body_markers("... protocol (Smith et al.,\n2020) is key.", _result())
    assert "\\cite{smith2020}" in tex


def test_out_of_range_numeric_marker_warns_and_is_left():
    tex, warnings = link_body_markers("Bad {[}99{]} marker.", _result())
    assert "{[}99{]}" in tex  # untouched
    assert "\\cite" not in tex
    assert any("99" in w for w in warnings)


def test_non_numeric_bracket_is_not_touched_or_warned():
    tex, warnings = link_body_markers("An aside {[}see appendix{]} here.", _result())
    assert tex == "An aside {[}see appendix{]} here."
    assert warnings == []


def test_superscript_ordinal_not_treated_as_citation():
    # "1\textsuperscript{st}" style ordinal has non-numeric content -> ignored.
    tex, warnings = link_body_markers("the 1\\textsuperscript{st} case", _result())
    assert tex == "the 1\\textsuperscript{st} case"
    assert warnings == []


def test_unknown_author_year_with_known_year_warns():
    result = _result(author_year_keys={("smith", "2020"): ["smith2020"]})
    tex, warnings = link_body_markers("As per (Nobody et al., 2020) here.", result)
    assert "(Nobody et al., 2020)" in tex  # untouched
    assert any("Nobody" in w for w in warnings)


def test_unknown_author_year_with_unknown_year_is_silent():
    # A parenthetical with a year the bibliography never mentions is likely prose.
    result = _result(author_year_keys={("smith", "2020"): ["smith2020"]})
    tex, warnings = link_body_markers("Founded (Acme Corp, 1889) long ago.", result)
    assert "(Acme Corp, 1889)" in tex
    assert warnings == []


def test_no_reference_list_leaves_markers_untouched():
    result = PlaintextResult(has_reference_list=False)
    tex, warnings = link_body_markers("Marker {[}1{]}.", result)
    # keys_by_number empty -> unresolved, warned, left in place.
    assert "{[}1{]}" in tex
    assert warnings


# --------------------------------------------------------------------------- #
# false-positive marker classes (crystallographic indices, title-page
# affiliation superscripts) -- GAP 3
# --------------------------------------------------------------------------- #


def test_miller_index_bracket_not_treated_as_citation():
    tex, warnings = link_body_markers("Grown along the {[}001{]} direction.", _result())
    assert tex == "Grown along the {[}001{]} direction."
    assert warnings == []


def test_multiple_miller_index_brackets_not_treated_as_citations():
    tex, warnings = link_body_markers(
        "Facets along {[}110{]} and {[}111{]} were observed.", _result()
    )
    assert tex == "Facets along {[}110{]} and {[}111{]} were observed."
    assert warnings == []


def test_miller_index_list_marker_fully_excluded_not_partially_warned():
    # All three canonical directions typed together as one bracketed list --
    # must not partially resolve/warn against whatever numbers happen to be
    # in keys_by_number (1, 2, 3, 4, 8 in _result()).
    tex, warnings = link_body_markers("Directions {[}001,110,111{]} shown.", _result())
    assert tex == "Directions {[}001,110,111{]} shown."
    assert warnings == []


def test_title_page_affiliation_superscripts_unwarned_but_in_body_marker_still_links():
    tex = (
        "J. Smith\\textsuperscript{1}, A. Doe\\textsuperscript{2,3} here.\n"
        "\\textsuperscript{1}Dept. of Physics; \\textsuperscript{2}Dept. of Chemistry.\n"
        "\\section{Introduction}\n"
        "Prior work\\textsuperscript{1} established the baseline; a later result"
        "\\textsuperscript{99} went further.\n"
    )
    new_tex, warnings = link_body_markers(tex, _result())

    # Title-page affiliation superscripts (before the first heading): left
    # untouched, no warning raised for them at all.
    assert "\\textsuperscript{1}Dept. of Physics" in new_tex
    assert "\\textsuperscript{2}Dept. of Chemistry" in new_tex
    assert "A. Doe\\textsuperscript{2,3}" in new_tex

    # A genuine in-body marker (after the heading) still resolves normally.
    assert "Prior work\\cite{smith2020} established" in new_tex

    # A genuine in-body out-of-range marker still warns -- the title-page
    # boundary must never suppress a real mismatch signal.
    assert "\\textsuperscript{99}" in new_tex
    assert len(warnings) == 1
    assert "99" in warnings[0]


def test_superscript_out_of_range_still_warns_with_no_title_page_content():
    tex = "\\section{Introduction}\nSee ref\\textsuperscript{99} here."
    new_tex, warnings = link_body_markers(tex, _result())
    assert "\\textsuperscript{99}" in new_tex
    assert any("99" in w for w in warnings)


def test_superscript_marker_resolves_normally_when_no_heading_present():
    # No sectioning command anywhere in the fragment (the common case for
    # these unit tests' bare snippets) -> no title-page boundary detected,
    # so nothing is suppressed; matches pre-GAP-3 behavior exactly.
    tex, warnings = link_body_markers("Body text with marker\\textsuperscript{1} only.", _result())
    assert "\\cite{smith2020}" in tex
    assert warnings == []


# --------------------------------------------------------------------------- #
# reference-section stripping
# --------------------------------------------------------------------------- #


def test_strip_removes_reference_section_to_eof():
    tex = (
        "\\section{Introduction}\\label{introduction}\n\n"
        "Body text.\n\n"
        "\\section{References}\\label{references}\n\n"
        "1. Smith, A. First. (2020).\n\n2. Jones, B. Second. (2019).\n"
    )
    stripped = strip_reference_section(tex, _result())
    assert "\\section{References}" not in stripped
    assert "Smith, A. First" not in stripped
    assert "\\section{Introduction}" in stripped
    assert "Body text." in stripped


def test_strip_removes_bold_reference_heading_not_promoted_to_section():
    # A Title-case "References" typed as a bold line is not ALL-CAPS, so
    # heading promotion leaves it as \textbf{References}; the fallback must
    # still strip the typed list from there to EOF.
    tex = (
        "\\section{Introduction}\\label{introduction}\n\n"
        "Body text mentioning the references shown in Fig. 1.\n\n"
        "\\textbf{References}\n\n"
        "1. Smith, A. First. (2020).\n\n2. Jones, B. Second. (2019).\n"
    )
    stripped = strip_reference_section(tex, _result())
    assert "\\textbf{References}" not in stripped
    assert "Smith, A. First" not in stripped
    # An in-sentence "references" earlier in the body must survive.
    assert "the references shown in Fig. 1" in stripped
    assert "Body text" in stripped


def test_strip_removes_bare_reference_heading_paragraph():
    tex = (
        "Body text.\n\n"
        "References\n\n"
        "1. Smith, A. First. (2020).\n"
    )
    stripped = strip_reference_section(tex, _result())
    assert stripped.rstrip().endswith("Body text.")


def test_strip_no_op_when_no_reference_heading():
    tex = "\\section{Introduction}\\label{introduction}\n\nBody only.\n"
    assert strip_reference_section(tex, _result()) == tex


def test_strip_no_op_when_result_has_no_reference_list():
    tex = "\\section{References}\\label{references}\n\n1. Smith. (2020).\n"
    result = PlaintextResult(has_reference_list=False)
    assert strip_reference_section(tex, result) == tex
