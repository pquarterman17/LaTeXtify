"""Generate ``paper.docx`` for example 01 (all-embedded).

The whole manuscript lives in one Word file: a title page, an abstract,
body sections, two *embedded* figures, in-text numeric citation markers,
and a typed reference list. Nothing sits beside it -- this is the
"I wrote everything in Word" happy path.

Figures are tiny solid-colour PNGs synthesized in-script via raw PNG chunk
encoding, so no binary assets are committed (the same trick the test suite
uses in ``tests/fixtures/make_figures.py``). Regenerate any time with::

    python make_manuscript.py
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_PATH = Path(__file__).with_name("paper.docx")


def _solid_png(color: tuple[int, int, int], size: int = 8) -> io.BytesIO:
    """A tiny solid-colour RGB PNG built by hand (no image library needed)."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0)  # 8-bit RGB
    scanlines = b"".join(b"\x00" + bytes(color) * size for _ in range(size))
    idat = zlib.compress(scanlines)
    return io.BytesIO(signature + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _superscript(paragraph, text: str) -> None:
    """Append a superscript run (used for affiliation markers on the author line)."""
    run = paragraph.add_run(text)
    run.font.superscript = True


def build(out_path: Path | str = OUT_PATH) -> Path:
    out_path = Path(out_path)
    document = Document()

    # --- Title page (LaTeXtify guesses metadata from this structure) ---------
    document.add_heading("Thermal Transport in Layered Antiferromagnets", level=0)  # Title style

    authors = document.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("Ada K. Researcher")
    _superscript(authors, "1")
    authors.add_run(", Ben T. Coauthor")
    _superscript(authors, "2")

    affil1 = document.add_paragraph()
    affil1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _superscript(affil1, "1")
    affil1.add_run("Department of Physics, Example University, Springfield, USA")

    affil2 = document.add_paragraph()
    affil2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _superscript(affil2, "2")
    affil2.add_run("National Laboratory for Materials, Metropolis, USA")

    document.add_paragraph("Abstract")  # exact text "Abstract" -> abstract detector
    document.add_paragraph(
        "We report long-distance magnon spin transport in a magnetic "
        "insulator and connect the measured signal to the emerging picture of "
        "magnon spintronics. The devices show a decay length consistent with "
        "diffusive magnon propagation at room temperature."
    )

    # --- Body ----------------------------------------------------------------
    # The in-text markers [1], [2] point at the typed reference list below.
    # These are REAL papers, so LaTeXtify's Crossref reconstruction resolves
    # them to clean BibTeX (with DOIs) when the machine is online; offline it
    # falls back to verify-flagged raw entries and still compiles.
    document.add_heading("Introduction", level=1)
    document.add_paragraph(
        "Magnons -- the quanta of spin waves -- can carry angular momentum "
        "over micrometre distances in magnetic insulators. The field of "
        "magnon spintronics has grown rapidly [1], and non-local devices have "
        "demonstrated room-temperature magnon spin transport in yttrium iron "
        "garnet [2]."
    )

    document.add_paragraph("Figure 1 shows the non-local device geometry.")
    document.add_picture(_solid_png((205, 60, 55)), width=Inches(2.4))
    document.add_paragraph(
        "Figure 1: Schematic of the non-local injector/detector device used to "
        "measure magnon spin transport.",
        style="Caption",
    )

    document.add_heading("Results and Discussion", level=1)
    document.add_paragraph(
        "The detected signal decays exponentially with injector-detector "
        "separation, giving a magnon relaxation length in agreement with the "
        "diffusive transport picture of Ref. [2]."
    )
    document.add_picture(_solid_png((55, 110, 200)), width=Inches(2.4))
    document.add_paragraph(
        "Figure 2: Non-local signal versus injector-detector separation.",
        style="Caption",
    )

    # --- Typed reference list (numeric markers above map onto these) ---------
    # Real, well-known references so Crossref matches them with high confidence.
    references = document.add_heading("References", level=1)
    references.runs[0].font.size = Pt(14)
    document.add_paragraph(
        "[1] A. V. Chumak, V. I. Vasyuchka, A. A. Serga, and B. Hillebrands, "
        "Magnon spintronics, Nature Physics 11, 453 (2015)."
    )
    document.add_paragraph(
        "[2] L. J. Cornelissen, J. Liu, R. A. Duine, J. Ben Youssef, and "
        "B. J. van Wees, Long-distance transport of magnon spin information in "
        "a magnetic insulator at room temperature, Nature Physics 11, 1022 "
        "(2015)."
    )

    document.save(out_path)
    return out_path


if __name__ == "__main__":
    written = build()
    print(f"wrote {written}")
