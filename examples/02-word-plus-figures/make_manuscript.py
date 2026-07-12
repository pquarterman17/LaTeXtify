"""Generate example 02: a manuscript whose real figures live *beside* the .docx.

A common real workflow: the Word file carries small, low-resolution draft
figures (whatever was pasted in while writing), but the publication-quality
files are kept separately. LaTeXtify overrides the embedded images with the
external ones, using two conventions that both sit next to ``paper.docx``:

    figures/fig1.png            <- folder convention: overrides figure 1
    figures.yaml  (2: ...)      <- explicit manifest: overrides figure 2
    figures/panels/signal.png       (the file the manifest points at)

Override precedence is manifest > folder > embedded, so a number listed in
``figures.yaml`` is taken from there even if a ``figures/figN.*`` file also
exists.

This script writes ALL of the above (the .docx, the external PNGs, and the
manifest) so the example is self-contained and commits no binaries.
Regenerate with::

    python make_manuscript.py
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

HERE = Path(__file__).parent
DOCX_PATH = HERE / "paper.docx"
FIGURES_DIR = HERE / "figures"
MANIFEST_PATH = HERE / "figures.yaml"


def _solid_png(color: tuple[int, int, int], size: int) -> bytes:
    """A solid-colour RGB PNG built by hand (no image library needed)."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0)  # 8-bit RGB
    scanlines = b"".join(b"\x00" + bytes(color) * size for _ in range(size))
    idat = zlib.compress(scanlines)
    return signature + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _build_docx() -> None:
    """The manuscript, with small GREY placeholder figures embedded inline."""
    document = Document()
    document.add_heading("Non-local Magnon Spin Transport in a Thin-Film Insulator", level=0)

    authors = document.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("Cora L. Devicemaker")
    sup = authors.add_run("1")
    sup.font.superscript = True

    affil = document.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker = affil.add_run("1")
    marker.font.superscript = True
    affil.add_run("Institute for Spintronics, Example University, Springfield, USA")

    document.add_paragraph("Abstract")
    document.add_paragraph(
        "We characterise non-local magnon spin transport in a thin-film "
        "magnetic insulator. Draft figures are embedded here at low "
        "resolution; the publication-quality versions are supplied alongside "
        "the document and substituted at conversion time."
    )

    document.add_heading("Device and Measurement", level=1)
    document.add_paragraph(
        "Figure 1 shows the device layout. The grey box below is only a "
        "low-resolution placeholder in the Word file."
    )
    document.add_picture(io.BytesIO(_solid_png((150, 150, 150), 8)), width=Inches(1.0))
    document.add_paragraph(
        "Figure 1: Optical micrograph of the non-local device.", style="Caption"
    )

    document.add_heading("Results", level=1)
    document.add_paragraph(
        "Figure 2 shows the non-local signal versus separation; again the "
        "embedded image is a placeholder, overridden via figures.yaml."
    )
    document.add_picture(io.BytesIO(_solid_png((150, 150, 150), 8)), width=Inches(1.0))
    document.add_paragraph(
        "Figure 2: Non-local voltage as a function of contact separation.",
        style="Caption",
    )

    document.save(DOCX_PATH)


def _build_external_figures() -> None:
    """Write the high-resolution external figures + the figures.yaml manifest."""
    (FIGURES_DIR / "panels").mkdir(parents=True, exist_ok=True)

    # Folder convention: figures/fig1.<ext> overrides figure 1 automatically.
    # A distinct colour (green) so you can SEE the override took effect.
    (FIGURES_DIR / "fig1.png").write_bytes(_solid_png((40, 160, 70), 64))

    # Manifest: figure 2 is taken from an arbitrary path via figures.yaml,
    # NOT the figures/ folder convention (this file isn't named fig2.*).
    (FIGURES_DIR / "panels" / "detector-signal.png").write_bytes(
        _solid_png((45, 90, 200), 64)
    )
    MANIFEST_PATH.write_text(
        "# figures.yaml -- maps a figure number to an explicit file path.\n"
        "# Paths are relative to this file. Manifest entries beat the\n"
        "# figures/figN.* folder convention, which beats the embedded image.\n"
        "2: figures/panels/detector-signal.png\n",
        encoding="utf-8",
    )


def build() -> Path:
    _build_docx()
    _build_external_figures()
    return DOCX_PATH


if __name__ == "__main__":
    written = build()
    print(f"wrote {written}")
    print(f"wrote {FIGURES_DIR / 'fig1.png'} (folder-convention override for figure 1)")
    print(f"wrote {FIGURES_DIR / 'panels' / 'detector-signal.png'} (manifest target for figure 2)")
    print(f"wrote {MANIFEST_PATH}")
